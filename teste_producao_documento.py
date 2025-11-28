from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

def set_cell_border(cell, **kwargs):
    """Define bordas para células da tabela"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'FFFFFF')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)

def set_cell_background(cell, color):
    """Define cor de fundo da célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Ler dados do Excel
df = pd.read_excel('exports/teste_integração.xlsx')

# Criar documento Word
doc = Document()

# Adicionar título
title = doc.add_heading('Resultados do Monitoramento de Metas Estratégicas 2024', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle1 = doc.add_paragraph('Relatório Técnico ao Comitê de Governança e Gestão Estratégica')
subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle2 = doc.add_paragraph('Presidência')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.color.rgb = RGBColor(255, 140, 0)

doc.add_paragraph()

# Criar tabela (assumindo 7 colunas)
table = doc.add_table(rows=1, cols=7)
table.style = 'Table Grid'

# Cabeçalho da tabela
header_cells = table.rows[0].cells
headers = ['INDICADOR', 'META', 'UNIDADE RESPONSÁVEL', 'POLARIDADE', 'RESULTADO APURADO', 'INICIATIVA(S)', 'IE']

for i, header in enumerate(headers):
    cell = header_cells[i]
    cell.text = header
    set_cell_background(cell, 'D2691E')  # Cor laranja
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = 1  # Centro vertical

# Adicionar dados do Excel
for index, row in df.iterrows():
    cells = table.add_row().cells
    for i, value in enumerate(row.values[:7]):  # Limitar a 7 colunas
        cells[i].text = str(value) if pd.notna(value) else ''
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Ajustar largura das colunas
widths = [Inches(2.5), Inches(2.5), Inches(1.2), Inches(1.0), Inches(1.0), Inches(1.5), Inches(0.8)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

# Salvar documento
doc.save('relatorio_metas_estrategicas.docx')
print("Documento gerado com sucesso!")