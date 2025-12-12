"""
Funções para criação e configuração do documento base
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION
from .config import Config


def criar_documento(superintendencia='Presidência'):
    """Cria documento Word base com primeira página em retrato"""
    doc = Document()
    
    # === PRIMEIRA SEÇÃO: RETRATO (para tabela histórica) ===
    section_retrato = doc.sections[0]
    
    # Orientação retrato
    section_retrato.orientation = WD_ORIENT.PORTRAIT
    
    # Tamanho do papel A4 em retrato
    section_retrato.page_width = Cm(21.0)
    section_retrato.page_height = Cm(29.7)
    
    # Margens
    section_retrato.top_margin = Cm(2.0)
    section_retrato.bottom_margin = Cm(1.5)
    section_retrato.left_margin = Cm(2.5)
    section_retrato.right_margin = Cm(2.5)
    
    # Medianiz (gutter)
    section_retrato.gutter = Cm(0)
    
    # Cabeçalho e rodapé - distâncias reduzidas para ficarem mais próximos das margens
    section_retrato.header_distance = Cm(0.5)
    section_retrato.footer_distance = Cm(1.27)
    
    # === CABEÇALHO DA PRIMEIRA PÁGINA (RETRATO) ===
    header_retrato = section_retrato.header
    
    # Limpar cabeçalho padrão
    for paragraph in list(header_retrato.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Linha 1: MONITORAMENTO DE METAS ESTRATÉGICAS - 2024 (negrito, tamanho 11, centralizado)
    p1 = header_retrato.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('MONITORAMENTO DE METAS ESTRATÉGICAS - 2024')
    run1.font.size = Pt(11)
    run1.font.bold = True
    run1.font.name = Config.FONTE_PADRAO
    p1.paragraph_format.space_after = Pt(0)
    
    # Linha 2: Relatório Técnico ao Comitê de Governança e Gestão Estratégica (normal, tamanho 11, centralizado)
    p2 = header_retrato.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Relatório Técnico ao Comitê de Governança e Gestão Estratégica')
    run2.font.size = Pt(11)
    run2.font.bold = False
    run2.font.name = Config.FONTE_PADRAO
    p2.paragraph_format.space_after = Pt(6)
    
    # === RODAPÉ DA PRIMEIRA PÁGINA (RETRATO) ===
    footer_retrato = section_retrato.footer
    
    # Limpar rodapé padrão
    for paragraph in list(footer_retrato.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Linha 1: Assessoria Técnica e Jurídica ao Planejamento e à Gestão Institucional - ASPLAG
    p_footer1 = footer_retrato.add_paragraph()
    p_footer1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_footer1 = p_footer1.add_run('Assessoria Técnica e Jurídica ao Planejamento e à Gestão Institucional - ASPLAG')
    run_footer1.font.size = Pt(9)
    run_footer1.font.name = Config.FONTE_PADRAO
    p_footer1.paragraph_format.space_after = Pt(0)
    
    # Linha 2: Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional - DEPLAG
    p_footer2 = footer_retrato.add_paragraph()
    p_footer2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_footer2 = p_footer2.add_run('Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional - DEPLAG')
    run_footer2.font.size = Pt(9)
    run_footer2.font.name = Config.FONTE_PADRAO
    p_footer2.paragraph_format.space_after = Pt(0)
    
    # === RETORNAR DOCUMENTO (segunda seção paisagem será criada depois) ===
    return doc


def criar_secao_paisagem_inicial(doc, superintendencia='Presidência'):
    """Cria segunda seção em paisagem para as tabelas de metas"""
    
    # Adicionar quebra de seção para paisagem
    section_paisagem = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Configurar orientação paisagem
    section_paisagem.orientation = WD_ORIENT.LANDSCAPE
    
    # Tamanho do papel A4 em paisagem
    section_paisagem.page_width = Cm(29.7)
    section_paisagem.page_height = Cm(21.0)
    
    # Margens
    section_paisagem.top_margin = Cm(2.5)
    section_paisagem.bottom_margin = Cm(2.5)
    section_paisagem.left_margin = Cm(2.0)
    section_paisagem.right_margin = Cm(2.0)
    
    # Medianiz (gutter)
    section_paisagem.gutter = Cm(0)
    
    # Cabeçalho e rodapé
    section_paisagem.header_distance = Cm(0.5)
    section_paisagem.footer_distance = Cm(1.27)
    
    # Desvincular cabeçalho da seção anterior
    header = section_paisagem.header
    header.is_linked_to_previous = False
    
    # Limpar cabeçalho padrão
    for paragraph in list(header.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Título principal - negrito, 12pt, centralizado
    titulo = header.add_paragraph('Resultados do Monitoramento de Metas Estratégicas 2025')
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_format = titulo.runs[0]
    titulo_format.font.size = Pt(12)
    titulo_format.font.bold = True
    titulo_format.font.name = Config.FONTE_PADRAO
    titulo.paragraph_format.space_after = Pt(0)
    titulo.paragraph_format.space_before = Pt(0)
    
    # Subtítulo - normal, 11pt, centralizado
    subtitulo = header.add_paragraph('Relatório Técnico ao Comitê de Governança e Gestão Estratégica')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_format = subtitulo.runs[0]
    subtitulo_format.font.size = Pt(11)
    subtitulo_format.font.name = Config.FONTE_PADRAO
    subtitulo.paragraph_format.space_after = Pt(0)
    
    # Órgão - negrito, 12pt, laranja, centralizado (DINÂMICO)
    orgao = header.add_paragraph(superintendencia)
    orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    orgao_format = orgao.runs[0]
    orgao_format.font.size = Pt(12)
    orgao_format.font.bold = True
    orgao_format.font.name = Config.FONTE_PADRAO
    orgao_format.font.color.rgb = RGBColor(227, 108, 10)
    orgao.paragraph_format.space_after = Pt(6)
    
    # Adicionar conteúdo no rodapé
    footer = section_paisagem.footer
    footer.is_linked_to_previous = False
    
    # Limpar rodapé padrão
    for paragraph in list(footer.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Criar tabela de 1 linha x 2 colunas para rodapé compacto
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    table = footer.add_table(rows=1, cols=2, width=Cm(50))
    table.autofit = False
    table.allow_autofit = False
    
    # Aplicar recuo negativo para estender além das margens
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '-582')  # -1.5cm em twips (1cm = 567 twips)
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Forçar layout fixo
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    # Configurar larguras das colunas em twips (1cm = 567 twips)
    table.columns[0].width = Cm(25.6)  # ASPLAG e DEPLAG - aumentado para paisagem
    table.columns[1].width = Cm(1.5)   # Número página
    
    # Configurar células
    row = table.rows[0]
    
    # Célula 1: ASPLAG e DEPLAG (esquerda)
    cell1 = row.cells[0]
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Definir largura explícita da célula 1: 24cm em twips
    tcPr1 = cell1._element.get_or_add_tcPr()
    tcW1 = OxmlElement('w:tcW')
    tcW1.set(qn('w:w'), str(int(25.6 * 567)))  # 24cm em twips
    tcW1.set(qn('w:type'), 'dxa')
    tcPr1.append(tcW1)
    # Desabilitar quebra automática
    noWrap = OxmlElement('w:noWrap')
    tcPr1.append(noWrap)
    
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p1.add_run('Assessoria Técnica e Jurídica ao Planejamento e à Gestão Institucional - ASPLAG')
    run1.font.size = Pt(9)
    run1.font.name = Config.FONTE_PADRAO
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    # Segunda linha - DEPLAG
    p1_linha2 = cell1.add_paragraph()
    p1_linha2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1_linha2 = p1_linha2.add_run('Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional - DEPLAG')
    run1_linha2.font.size = Pt(9)
    run1_linha2.font.name = Config.FONTE_PADRAO
    p1_linha2.paragraph_format.space_before = Pt(0)
    p1_linha2.paragraph_format.space_after = Pt(0)
    
    # Célula 2: Número da página (direita)
    cell2 = row.cells[1]
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Definir largura explícita da célula 2: 1.5cm em twips
    tcPr2 = cell2._element.get_or_add_tcPr()
    tcW2 = OxmlElement('w:tcW')
    tcW2.set(qn('w:w'), str(int(1.5 * 567)))  # 1.5cm em twips
    tcW2.set(qn('w:type'), 'dxa')
    tcPr2.append(tcW2)
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run()
    run2.font.size = Pt(12)
    run2.font.name = Config.FONTE_PADRAO
    # Adicionar campo PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    
    # Remover bordas laterais e inferior, manter apenas borda superior
    for i, cell in enumerate(row.cells):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        # Borda superior (linha fina)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')  # 0.5pt
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        
        # Sem bordas laterais e inferior
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'none')
        tcBorders.append(left)
        
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'none')
        tcBorders.append(right)
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'none')
        tcBorders.append(bottom)
        
        tcPr.append(tcBorders)
    
    return section_paisagem
