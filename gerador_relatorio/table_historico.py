"""
Módulo para tabelas históricas
Contém funções para adicionar tabelas de histórico de metas aprovadas e por macrodesafio
"""

import pandas as pd
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from .config import Config


def adicionar_tabela_historica(doc):
    """
    Adiciona tabela histórica de metas aprovadas (4 anos mais recentes).
    Formato conforme imagem de referência com cabeçalho laranja.
    """
    from base_dados_fixos import atualizar_historico_com_ano_atual
    
    # Obter dados históricos (4 anos mais recentes)
    dados = atualizar_historico_com_ano_atual()
    historico = dados['historico_recente']
    anos = dados['anos_recentes']
    variacao = dados['variacao']
    primeiro_ano = anos[0]
    ultimo_ano = anos[-1]
    
    # Adicionar parágrafo de espaçamento
    doc.add_paragraph()
    
    # Criar tabela (5 linhas x 6 colunas)
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'
    
    # === LINHA 1: TÍTULO (MERGED) ===
    titulo_cells = table.rows[0].cells
    merged_cell = titulo_cells[0].merge(titulo_cells[5])
    
    titulo_paragraph = merged_cell.paragraphs[0]
    titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_paragraph.add_run(
        'TOTAL DE METAS APROVADAS PARA COMPOSIÇÃO DO PLANEJAMENTO ESTRATÉGICO\nINSTITUCIONAL DO TJMG'
    )
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor(255, 255, 255)
    titulo_run.font.name = Config.FONTE_PADRAO
    
    shading = parse_xml(r'<w:shd {} w:fill="E36C0A"/>'.format(nsdecls('w')))
    merged_cell._element.get_or_add_tcPr().append(shading)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].height = Cm(0.7)
    
    # === LINHA 2: CABEÇALHO ===
    header_row = table.rows[1]
    headers = ['Ano'] + [str(ano) for ano in anos] + [f'Variação\n{primeiro_ano} - {ultimo_ano}']
    
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="FABF8F"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[1].height = Cm(0.7)
    
    # === LINHAS 3-5: DADOS ===
    linhas_dados = [
        ('Metas Nacionais', 2),
        ('Metas Institucionais', 3),
        ('Total', 4)
    ]
    
    for label, idx_linha in linhas_dados:
        row = table.rows[idx_linha]
        row.cells[0].text = label
        
        for idx, ano in enumerate(anos, start=1):
            if label == 'Metas Nacionais':
                row.cells[idx].text = str(historico[ano]['nacionais'])
            elif label == 'Metas Institucionais':
                row.cells[idx].text = str(historico[ano]['institucionais'])
            else:  # Total
                row.cells[idx].text = str(historico[ano]['total'])
        
        # Variação
        if label == 'Metas Nacionais':
            row.cells[5].text = f"{variacao['nacionais']:+}%"
        elif label == 'Metas Institucionais':
            row.cells[5].text = f"{variacao['institucionais']:+}%"
        else:
            row.cells[5].text = f"{variacao['total']:+}%"
        
        # Formatação
        for idx, cell in enumerate(row.cells):
            paragraph = cell.paragraphs[0]
            
            if idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.bold = (label == 'Total')
                run.font.name = Config.FONTE_PADRAO
            
            # Linha Total: fundo cinza escuro e texto branco
            if label == 'Total':
                shading = parse_xml(r'<w:shd {} w:fill="404040"/>'.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(shading)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        table.rows[idx_linha].height = Cm(0.7 if label == 'Total' else 1.0)
    
    # Ajustar largura das colunas
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        for i in range(1, 6):
            row.cells[i].width = Cm(2.5)
    
    # Nota de rodapé
    nota = doc.add_paragraph()
    nota_run = nota.add_run(
        '*Destaca-se que a redução no número de metas nacionais no período decorreu de revisões propostas pelo '
        'Conselho Nacional de Justiça – CNJ, e aprovada pelos Tribunais Estaduais, em relação a seu compromisso com '
        'o Sistema de Justiça e a prestação jurisdicional.'
    )
    nota_run.font.size = Pt(8)
    nota_run.font.italic = True
    nota_run.font.name = Config.FONTE_PADRAO
    nota.paragraph_format.space_before = Pt(6)
    
    doc.add_paragraph()


def adicionar_tabela_macrodesafio(doc):
    """
    Adiciona tabela histórica de metas por macrodesafio (4 anos mais recentes).
    """
    from base_dados_fixos import atualizar_historico_macrodesafio_com_ano_atual
    
    dados = atualizar_historico_macrodesafio_com_ano_atual()
    historico = dados['historico_recente']
    anos = dados['anos_recentes']
    macrodesafios = dados['macrodesafios']
    
    doc.add_paragraph()
    
    num_linhas = 2 + len(macrodesafios) + 1
    num_colunas = 1 + len(anos)
    
    table = doc.add_table(rows=num_linhas, cols=num_colunas)
    table.style = 'Table Grid'
    
    # Título
    titulo_cells = table.rows[0].cells
    merged_cell = titulo_cells[0].merge(titulo_cells[-1])
    
    titulo_paragraph = merged_cell.paragraphs[0]
    titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_paragraph.add_run(
        f'TOTAL DE METAS POR MACRODESAFIO – COMPARATIVO {anos[0]} A {anos[-1]}'
    )
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor(255, 255, 255)
    titulo_run.font.name = Config.FONTE_PADRAO
    
    shading = parse_xml(r'<w:shd {} w:fill="E36C0A"/>'.format(nsdecls('w')))
    merged_cell._element.get_or_add_tcPr().append(shading)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].height = Cm(0.8)
    
    # Cabeçalho
    header_row = table.rows[1]
    headers = ['Macrodesafio'] + [str(ano) for ano in anos]
    
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="FABF8F"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[1].height = Cm(0.8)
    
    # Linhas de macrodesafios
    total_por_ano = {ano: 0 for ano in anos}
    
    for idx_macro, macrodesafio in enumerate(macrodesafios, start=2):
        row = table.rows[idx_macro]
        idx_relativo = idx_macro - 2
        cor_fundo = "FBD4B4" if idx_relativo % 2 == 0 else "FFFFFF"
        
        row.cells[0].text = macrodesafio
        paragraph = row.cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[0]._element.get_or_add_tcPr().append(shading)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        for idx_ano, ano in enumerate(anos, start=1):
            valor = historico[ano].get(macrodesafio, 0)
            total_por_ano[ano] += valor
            
            row.cells[idx_ano].text = str(valor)
            paragraph = row.cells[idx_ano].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            is_ultimo_ano = (idx_ano == len(anos))
            
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.bold = is_ultimo_ano
                run.font.name = Config.FONTE_PADRAO
            
            shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
            row.cells[idx_ano]._element.get_or_add_tcPr().append(shading)
            row.cells[idx_ano].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        table.rows[idx_macro].height = Cm(0.7)
    
    # Linha Total
    row_total_idx = num_linhas - 1
    row_total = table.rows[row_total_idx]
    
    row_total.cells[0].text = 'Total Geral'
    paragraph = row_total.cells[0].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
    
    shading = parse_xml(r'<w:shd {} w:fill="404040"/>'.format(nsdecls('w')))
    row_total.cells[0]._element.get_or_add_tcPr().append(shading)
    
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    row_total.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    for idx_ano, ano in enumerate(anos, start=1):
        row_total.cells[idx_ano].text = str(total_por_ano[ano])
        paragraph = row_total.cells[idx_ano].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
            run.font.color.rgb = RGBColor(255, 255, 255)
        
        shading = parse_xml(r'<w:shd {} w:fill="404040"/>'.format(nsdecls('w')))
        row_total.cells[idx_ano]._element.get_or_add_tcPr().append(shading)
        
        row_total.cells[idx_ano].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[row_total_idx].height = Cm(0.8)
    
    # Ajustar larguras
    for row in table.rows:
        row.cells[0].width = Cm(12.0)
        for i in range(1, num_colunas):
            row.cells[i].width = Cm(1.94)
    
    # Aplicar recuo negativo
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr {}/> '.format(nsdecls('w')))
        tbl.insert(0, tblPr)
    
    tblInd = parse_xml(r'<w:tblInd {} w:w="-397" w:type="dxa"/>'.format(nsdecls('w')))
    tblPr.append(tblInd)
    
    doc.add_paragraph()
