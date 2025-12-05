"""
Módulo para tabelas e seções de superintendências
Contém funções para adicionar seções de macrodesafios e tabelas de indicadores
"""

import pandas as pd
import re
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .config import Config
from .formatters import formatar_valor
from .styles import set_cell_background, set_cell_border


def adicionar_nova_secao_superintendencia(doc, superintendencia, primeira=False):
    """Adiciona nova seção com cabeçalho personalizado para superintendência"""
    
    if not primeira:
        # Adicionar quebra de seção (cria nova seção automaticamente)
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        # Usar a última seção existente
        new_section = doc.sections[-1]
    
    # Configurar a nova seção
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.top_margin = Cm(2.5)
    new_section.left_margin = Cm(2.0)
    new_section.bottom_margin = Cm(2.5)
    new_section.right_margin = Cm(2.0)
    new_section.gutter = Cm(0)
    new_section.header_distance = Cm(1.05)
    new_section.footer_distance = Cm(1.05)
    
    # IMPORTANTE: Desvincular cabeçalho e rodapé da seção anterior
    new_section.different_first_page_header_footer = False
    
    # Configurar cabeçalho da nova seção
    header = new_section.header
    
    # Desvincular do cabeçalho anterior
    header.is_linked_to_previous = False
    
    # Limpar cabeçalho padrão
    for paragraph in list(header.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Título principal
    titulo = header.add_paragraph('Resultados do Monitoramento de Metas Estratégicas 2025')
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_format = titulo.runs[0]
    titulo_format.font.size = Pt(12)
    titulo_format.font.bold = True
    titulo_format.font.name = Config.FONTE_PADRAO
    titulo.paragraph_format.space_after = Pt(0)
    titulo.paragraph_format.space_before = Pt(0)
    
    # Subtítulo
    subtitulo = header.add_paragraph('Relatório Técnico ao Comitê de Governança e Gestão Estratégica')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_format = subtitulo.runs[0]
    subtitulo_format.font.size = Pt(11)
    subtitulo_format.font.name = Config.FONTE_PADRAO
    subtitulo.paragraph_format.space_after = Pt(0)
    
    # Superintendência (DINÂMICO)
    orgao = header.add_paragraph(superintendencia)
    orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    orgao_format = orgao.runs[0]
    orgao_format.font.size = Pt(12)
    orgao_format.font.bold = True
    orgao_format.font.name = Config.FONTE_PADRAO
    orgao_format.font.color.rgb = RGBColor(227, 108, 10)
    orgao.paragraph_format.space_after = Pt(6)
    
    # Configurar rodapé da nova seção
    footer = new_section.footer
    
    # Desvincular do rodapé anterior
    footer.is_linked_to_previous = False
    
    # Limpar rodapé padrão
    for paragraph in list(footer.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Criar tabela de 1 linha x 2 colunas para rodapé compacto
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    table = footer.add_table(rows=1, cols=2, width=Cm(25))
    table.autofit = False
    
    # Configurar larguras das colunas (paisagem tem mais espaço)
    table.columns[0].width = Cm(22)  # ASPLAG e DEPLAG
    table.columns[1].width = Cm(3)   # Número página
    
    # Configurar células
    row = table.rows[0]
    
    # Célula 1: ASPLAG e DEPLAG (esquerda)
    cell1 = row.cells[0]
    cell1.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p1.add_run('Assessoria Técnica e Jurídica ao Planejamento e à Gestão Institucional - ASPLAG / Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional - DEPLAG')
    run1.font.size = Pt(9)
    run1.font.name = Config.FONTE_PADRAO
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    
    # Célula 2: Número da página (direita)
    cell2 = row.cells[1]
    cell2.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run()
    run2.font.size = Pt(9)
    run2.font.name = Config.FONTE_PADRAO
    # Adicionar campo PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    
    # Remover bordas laterais e inferior, manter apenas borda superior
    for cell in row.cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        # Borda superior (linha fina)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')  # 0.5pt
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        # Sem bordas laterais e inferior
        for border in ['left', 'bottom', 'right']:
            elem = OxmlElement(f'w:{border}')
            elem.set(qn('w:val'), 'none')
            tcBorders.append(elem)
        tcPr.append(tcBorders)
        
        # Remover padding interno para altura mínima
        tcMar = OxmlElement('w:tcMar')
        for margin in ['top', 'bottom', 'left', 'right']:
            mar = OxmlElement(f'w:{margin}')
            mar.set(qn('w:w'), '0')
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        tcPr.append(tcMar)


def adicionar_secao_macrodesafio(doc, macrodesafio, df_grupo, primeira_secao=False):
    """Adiciona seção de um Macrodesafio"""
    
    # Quebra de página entre macrodesafios (exceto no primeiro)
    if not primeira_secao:
        doc.add_page_break()
    
    def adicionar_cabecalho_macro():
        """Função auxiliar para adicionar cabeçalho do macrodesafio"""
        # Título do Macrodesafio usando tabela para controlar largura
        titulo_tabela = doc.add_table(rows=1, cols=1)
        titulo_tabela.style = None
        titulo_cell = titulo_tabela.rows[0].cells[0]
        titulo_cell.text = macrodesafio.upper()
        
        # Formatação do título
        for paragraph in titulo_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(10)  # Espaço antes
            paragraph.paragraph_format.space_after = Pt(10)   # Espaço depois
            paragraph.paragraph_format.line_spacing = 1.0     # Espaçamento entre linhas
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = Config.FONTE_PADRAO
                run.font.color.rgb = RGBColor(*Config.CORES['TEXTO_BRANCO'])
                run.font.bold = True
        
        # Cor de fundo e alinhamento vertical
        set_cell_background(titulo_cell, Config.CORES['CABECALHO_MACRO'])
        titulo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Definir altura da célula: 1,03cm
        titulo_tabela.rows[0].height = Cm(1.03)
        
        # Adicionar borda laranja na esquerda (1pt de espessura)
        set_cell_border(titulo_cell, 
                       top=False, 
                       bottom=False, 
                       left={'color': (227, 108, 10), 'size': 1}, 
                       right=False)
        
        # Definir largura igual à soma das colunas da tabela de dados
        # Largura total: 28,5cm (de -1cm até 27,5cm na régua)
        largura_total_tabela = Cm(28.5)
        titulo_tabela.rows[0].cells[0].width = largura_total_tabela
        
        # Adicionar recuo negativo de 1cm
        tbl = titulo_tabela._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '-567')  # -1cm em twips (1cm = 567 twips)
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        # Marcar que já foi executado uma vez
        adicionar_cabecalho_macro._ja_executado = True
    
    # Adicionar cabeçalho inicial
    adicionar_cabecalho_macro()
    
    # Processar cada indicador individualmente com controle de quebra
    precisa_cabecalho = False
    incluir_subcabecalho = True  # Primeiro sempre tem subcabeçalho
    
    for idx, (_, row) in enumerate(df_grupo.iterrows()):
        # Se precisa cabeçalho (após quebra de página), adicionar
        if precisa_cabecalho:
            doc.add_page_break()
            adicionar_cabecalho_macro()
            incluir_subcabecalho = True  # Após quebra, sempre incluir subcabeçalho
            precisa_cabecalho = False
        
        # Criar tabela para este indicador
        largura_total = adicionar_tabela_indicador(doc, row, incluir_cabecalho=incluir_subcabecalho)
        incluir_subcabecalho = False  # Próximos não terão subcabeçalho (a menos que haja quebra)
        
        # Adicionar informação complementar se existir
        info_complementar = row.get(Config.COLUNAS['INFO_COMPLEMENTAR'], '')
        tem_situacao = pd.notna(info_complementar) and str(info_complementar).strip() != ''
        
        if tem_situacao:
            texto_situacao = str(info_complementar)
            
            # Criar tabela de célula única para situação
            situacao_tabela = doc.add_table(rows=1, cols=1)
            situacao_tabela.style = None
            situacao_cell = situacao_tabela.rows[0].cells[0]
            
            # Processar texto e adicionar formatação
            # Dividir o texto em partes: antes dos marcadores e itens com marcadores
            partes = re.split(r'(•)', texto_situacao)
            
            primeiro_paragrafo = True
            paragrafo_atual = None
            texto_acumulado = ""
            
            for i, parte in enumerate(partes):
                if parte == '•':
                    # Finalizar parágrafo anterior se houver texto acumulado
                    if texto_acumulado.strip():
                        if paragrafo_atual is None:
                            paragrafo_atual = situacao_cell.paragraphs[0] if primeiro_paragrafo and len(situacao_cell.paragraphs) > 0 else situacao_cell.add_paragraph()
                            primeiro_paragrafo = False
                        
                        if paragrafo_atual.text == '':
                            run_bold = paragrafo_atual.add_run('Situação: ')
                            run_bold.font.bold = True
                            run_bold.font.size = Pt(10)
                            run_bold.font.name = Config.FONTE_PADRAO
                        
                        run = paragrafo_atual.add_run(texto_acumulado)
                        run.font.size = Pt(10)
                        run.font.name = Config.FONTE_PADRAO
                        
                        paragrafo_atual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragrafo_atual.paragraph_format.space_before = Pt(0)
                        paragrafo_atual.paragraph_format.space_after = Pt(0)
                        
                        texto_acumulado = ""
                        paragrafo_atual = None
                    
                    # Iniciar novo parágrafo com marcador
                    paragrafo_atual = situacao_cell.add_paragraph()
                    primeiro_paragrafo = False
                    
                    run_marcador = paragrafo_atual.add_run('•')
                    run_marcador.font.size = Pt(10)
                    run_marcador.font.name = Config.FONTE_PADRAO
                    
                    paragrafo_atual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragrafo_atual.paragraph_format.space_before = Pt(0)
                    paragrafo_atual.paragraph_format.space_after = Pt(0)
                else:
                    texto_acumulado += parte
            
            # Adicionar último texto acumulado
            if texto_acumulado.strip():
                if paragrafo_atual is None:
                    if len(situacao_cell.paragraphs) > 0 and primeiro_paragrafo:
                        paragrafo_atual = situacao_cell.paragraphs[0]
                    else:
                        paragrafo_atual = situacao_cell.add_paragraph()
                
                # Adicionar "Situação:" apenas se for o primeiro parágrafo e estiver vazio
                if len(situacao_cell.paragraphs) == 1 and paragrafo_atual.text == '':
                    run_bold = paragrafo_atual.add_run('Situação: ')
                    run_bold.font.bold = True
                    run_bold.font.size = Pt(10)
                    run_bold.font.name = Config.FONTE_PADRAO
                
                run = paragrafo_atual.add_run(texto_acumulado)
                run.font.size = Pt(10)
                run.font.name = Config.FONTE_PADRAO
                
                paragrafo_atual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragrafo_atual.paragraph_format.space_before = Pt(0)
                paragrafo_atual.paragraph_format.space_after = Pt(0)
            
            # Formatação da célula
            situacao_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Definir mesma largura da tabela principal
            situacao_tabela.rows[0].cells[0].width = largura_total
            
            # Adicionar bordas
            set_cell_border(situacao_cell, top=True, bottom=True, left=True, right=True)
            
            # Adicionar recuo negativo de 1cm à tabela de situação
            tbl_situacao = situacao_tabela._element
            tblPr_situacao = tbl_situacao.tblPr
            if tblPr_situacao is None:
                tblPr_situacao = OxmlElement('w:tblPr')
                tbl_situacao.insert(0, tblPr_situacao)
            
            tblInd_situacao = OxmlElement('w:tblInd')
            tblInd_situacao.set(qn('w:w'), '-567')  # -1cm em twips
            tblInd_situacao.set(qn('w:type'), 'dxa')
            tblPr_situacao.append(tblInd_situacao)
            
            # Permitir que a tabela de situação quebre entre páginas se necessário
            for row in situacao_tabela.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    # Remover restrição cantSplit se existir
                    cantSplit_elements = tcPr.findall(qn('w:cantSplit'))
                    for cs in cantSplit_elements:
                        tcPr.remove(cs)
        
        # Verificar se próximo indicador precisa de nova página
        if tem_situacao and len(str(info_complementar)) > 800:
            precisa_cabecalho = True


def adicionar_tabela_indicador(doc, row, incluir_cabecalho=True):
    """Adiciona tabela com um único indicador e retorna a largura total"""
    
    # Criar tabela (cabeçalho opcional + 1 linha de dados)
    num_linhas = 2 if incluir_cabecalho else 1
    num_colunas = 7  # Incluindo coluna de indicador visual
    tabela = doc.add_table(rows=num_linhas, cols=num_colunas)
    tabela.style = None  # Remover estilo padrão para controle total
    
    # Cabeçalhos (se necessário)
    if incluir_cabecalho:
        headers = ['INDICADOR', 'META', 'UNIDADE RESPONSÁVEL', 'POLARIDADE', 'RESULTADO APURADO', 'INICIATIVA(S)']
        header_cells = tabela.rows[0].cells
        
        # Mesclar células 4 e 5 (RESULTADO APURADO e ●)
        header_cells[4].merge(header_cells[5])
        
        for i, header in enumerate(headers):
            # Pular índice 5 pois foi mesclado com 4
            if i >= 5:
                cell_index = i + 1
            else:
                cell_index = i
            
            cell = header_cells[cell_index]
            cell.text = header
            
            # Formatação do cabeçalho
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = Config.FONTE_PADRAO
                    run.font.color.rgb = RGBColor(*Config.CORES['TEXTO_PRETO'])
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Cor de fundo
            set_cell_background(cell, Config.CORES['CABECALHO_TABELA'])
            
            # Adicionar bordas brancas
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
        
        # Definir altura da linha do cabeçalho: 1,03cm
        tabela.rows[0].height = Cm(1.03)
    
    # Preencher dados (última linha da tabela)
    linha_dados = 1 if incluir_cabecalho else 0
    cells = tabela.rows[linha_dados].cells
    
    # Dados
    dados = [
        formatar_valor(row.get(Config.COLUNAS['INDICADOR'], '-')),
        formatar_valor(row.get(Config.COLUNAS['METAKEY'], '-')),  # MetaKey contém o texto completo
        formatar_valor(row.get(Config.COLUNAS['UNIDADE_GESTORA'], '-')),
        formatar_valor(row.get(Config.COLUNAS['POLARIDADE'], '-')),
        formatar_valor(row.get(Config.COLUNAS['VALOR_APURADO'], '-')),
        '●',  # Indicador visual
        formatar_valor(row.get(Config.COLUNAS['INICIATIVA'], '-'))
    ]
    
    # Determinar cor do indicador visual baseado no atingimento da meta
    valor_apurado = row.get(Config.COLUNAS['VALOR_APURADO'], None)
    valor_meta = row.get(Config.COLUNAS['VALOR_META'], None)
    
    # Verde se atingiu a meta, vermelho caso contrário
    cor_indicador = RGBColor(0, 255, 0)  # Verde (padrão)
    if pd.notna(valor_apurado) and pd.notna(valor_meta):
        try:
            # Converter valor apurado (pode estar como string com vírgula)
            if isinstance(valor_apurado, str):
                val_apurado_float = float(valor_apurado.replace('.', '').replace(',', '.'))
            else:
                val_apurado_float = float(valor_apurado)
            
            # Converter valor da meta
            val_meta_float = float(valor_meta)
            
            if val_apurado_float < val_meta_float:
                cor_indicador = RGBColor(255, 0, 0)  # Vermelho
        except (ValueError, TypeError):
            pass  # Mantém verde se não conseguir converter
    
    for i, dado in enumerate(dados):
        cell = cells[i]
        
        # Tratamento especial para coluna Indicador (i == 0)
        if i == 0:
            cell.text = ''
            paragraph = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
            
            # Separar número do texto (ex: "1.14 - Texto" ou "1.14 Texto")
            match = re.match(r'^(\d+\.?\d*)\s*[-–]?\s*(.*)$', dado)
            
            if match:
                numero = match.group(1)
                texto_resto = match.group(2)
                
                # Adicionar número em negrito
                run_negrito = paragraph.add_run(numero)
                run_negrito.font.bold = True
                run_negrito.font.size = Pt(Config.TAMANHO_TABELA)
                run_negrito.font.name = Config.FONTE_PADRAO
                
                # Adicionar o resto do texto normalmente
                if texto_resto:
                    run_normal = paragraph.add_run(' - ' + texto_resto)
                    run_normal.font.size = Pt(Config.TAMANHO_TABELA)
                    run_normal.font.name = Config.FONTE_PADRAO
            else:
                # Se não encontrar padrão, adicionar texto normal
                run = paragraph.add_run(dado)
                run.font.size = Pt(Config.TAMANHO_TABELA)
                run.font.name = Config.FONTE_PADRAO
            
            # Alinhamento
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        
        # Tratamento especial para coluna Meta (i == 1)
        elif i == 1:
            cell.text = ''
            paragraph = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
            
            # Separar código da meta do texto (ex: "TJMG 111 - Texto" ou "TJMG 111")
            match = re.match(r'^([A-Z]+\s+\d+)\s*[-–]?\s*(.*)$', dado)
            
            if match:
                codigo = match.group(1)
                texto_resto = match.group(2)
                
                # Adicionar código em negrito
                run_negrito = paragraph.add_run(codigo)
                run_negrito.font.bold = True
                run_negrito.font.size = Pt(Config.TAMANHO_TABELA)
                run_negrito.font.name = Config.FONTE_PADRAO
                
                # Adicionar o resto do texto normalmente
                if texto_resto:
                    run_normal = paragraph.add_run(' - ' + texto_resto)
                    run_normal.font.size = Pt(Config.TAMANHO_TABELA)
                    run_normal.font.name = Config.FONTE_PADRAO
            else:
                # Se não encontrar padrão, adicionar texto normal
                run = paragraph.add_run(dado)
                run.font.size = Pt(Config.TAMANHO_TABELA)
                run.font.name = Config.FONTE_PADRAO
            
            # Alinhamento
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        
        else:
            # Outras colunas: formatação normal
            cell.text = dado
        
        # Formatação do texto (para colunas que não são Indicador nem Meta)
        if i not in [0, 1]:
            for paragraph in cell.paragraphs:
                # Alinhamento - centralizar todos os dados
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Remover espaçamentos
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            
            for run in paragraph.runs:
                run.font.size = Pt(Config.TAMANHO_TABELA)
                run.font.name = Config.FONTE_PADRAO
                
                # Destacar resultado apurado
                if i == 4:  # Coluna de Resultado Apurado
                    run.font.bold = True
                
                # Cor do indicador visual
                if i == 5:  # Coluna ●
                    run.font.color.rgb = cor_indicador
                    run.font.size = Pt(16)
        
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Adicionar bordas brancas
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)
    
    # Ajustar largura das colunas (em centímetros)
    larguras = [
        None,        # INDICADOR - calculado depois
        Cm(9.09),    # META
        Cm(2.28),    # UNIDADE RESPONSÁVEL
        Cm(2.1),     # POLARIDADE
        Cm(1.97),    # RESULTADO APURADO
        Cm(0.72),    # ●
        Cm(2.18)     # INICIATIVA(S)
    ]
    
    # Calcular largura do Indicador baseado no espaço restante
    largura_total = Cm(28.5)
    largura_outras = sum([l for l in larguras if l is not None])
    largura_indicador = largura_total - largura_outras
    larguras[0] = max(largura_indicador, Cm(1.3))
    
    for row_table in tabela.rows:
        for idx, width in enumerate(larguras):
            if width:
                row_table.cells[idx].width = width
    
    # Calcular largura total da tabela para retornar
    largura_total_tabela = sum([l for l in larguras if l is not None])
    
    # Configurar propriedades da tabela
    tbl = tabela._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Adicionar recuo negativo de 0,7cm
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '-397')  # -0,7cm em twips
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Permitir que a tabela quebre entre linhas se necessário
    tblPrEx = OxmlElement('w:tblPrEx')
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), '0')
    tblPrEx.append(cantSplit)
    
    return largura_total_tabela
