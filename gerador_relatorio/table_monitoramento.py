"""
Módulo para tabela de resultado do monitoramento
Contém função para adicionar tabela com distribuição por faixa de cumprimento
"""

import pandas as pd
import re
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from .config import Config


def adicionar_tabela_resultado_monitoramento(doc):
    """
    Adiciona tabela de resultado do monitoramento de metas.
    Mostra distribuição por faixa de cumprimento.
    Une dados de metas institucionais (TJMG) e metas nacionais (CNJ).
    """
    from base_dados_fixos import calcular_cumprimento_metas_cnj
    
    print("\n📊 Calculando resultado do monitoramento...")
    
    # === CALCULAR CUMPRIMENTO METAS CNJ ===
    metas_cnj_cumprimento = calcular_cumprimento_metas_cnj()
    
    # === CALCULAR CUMPRIMENTO METAS INSTITUCIONAIS ===
    try:
        df_tjmg = pd.read_excel('exports/teste_integração.xlsx')
        
        # Converter percentual para float
        def converter_percentual(val):
            if pd.isna(val):
                return 0.0
            val_str = str(val).replace('%', '').replace(',', '.').strip()
            try:
                return float(val_str)
            except:
                return 0.0
        
        df_tjmg['Resultado_Float'] = df_tjmg['Valor Apurado'].apply(converter_percentual)
        
        # Para metas institucionais, usar valor apurado diretamente
        metas_tjmg_cumprimento = {}
        for _, row in df_tjmg.iterrows():
            meta_key = row['MetaKey']
            # Extrair código (TJMG XXX)
            match = re.search(r'^(TJMG\s+\d+)', meta_key)
            if match:
                codigo = match.group(1)
                cumprimento = row['Resultado_Float']
                metas_tjmg_cumprimento[codigo] = {'cumprimento': cumprimento, 'categorias': 1}
    
    except Exception as e:
        print(f"⚠️  Erro ao processar metas TJMG: {e}")
        metas_tjmg_cumprimento = {}
    
    # === UNIR TODOS OS CUMPRIMENTOS ===
    todos_cumprimentos = {}
    todos_cumprimentos.update(metas_cnj_cumprimento)
    todos_cumprimentos.update(metas_tjmg_cumprimento)
    
    # === CLASSIFICAR POR FAIXA ===
    maior_100 = 0
    entre_70_100 = 0
    abaixo_70 = 0
    sem_apuracao = 0
    
    for meta, dados in todos_cumprimentos.items():
        cumprimento = dados['cumprimento']
        
        if cumprimento >= 100:
            maior_100 += 1
        elif cumprimento >= 70:
            entre_70_100 += 1
        elif cumprimento > 0:
            abaixo_70 += 1
        else:
            sem_apuracao += 1
    
    total = len(todos_cumprimentos)
    
    print(f"✅ Processadas {total} metas ({len(metas_cnj_cumprimento)} CNJ + {len(metas_tjmg_cumprimento)} TJMG)")
    
    # Adicionar parágrafo de espaçamento
    doc.add_paragraph()
    
    # Criar tabela (linhas: título + 4 faixas + total) x 3 colunas (SEM cabeçalho)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    # === LINHA 1: TÍTULO (MERGED) ===
    titulo_cells = table.rows[0].cells
    merged_cell = titulo_cells[0].merge(titulo_cells[-1])
    
    titulo_paragraph = merged_cell.paragraphs[0]
    titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_paragraph.add_run('RESULTADO DO MONITORAMENTO DAS METAS ESTRATÉGICAS')
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor(255, 255, 255)
    titulo_run.font.name = Config.FONTE_PADRAO
    
    # Cor de fundo laranja
    shading = parse_xml(r'<w:shd {} w:fill="E36C0A"/>'.format(nsdecls('w')))
    merged_cell._element.get_or_add_tcPr().append(shading)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].height = Cm(0.8)
    
    # === LINHAS DE DADOS (começam na linha 1, índice ajustado) ===
    faixas = [
        ('Metas com resultado maior ou igual 100%', maior_100),
        ('Metas com resultado entre 70% e 100%', entre_70_100),
        ('Metas com resultado abaixo de 70%', abaixo_70),
        ('Metas sem apuração até outubro/2025', sem_apuracao)
    ]
    
    for idx_faixa, (descricao, quantidade) in enumerate(faixas, start=1):
        row = table.rows[idx_faixa]
        
        # Definir cor de fundo alternada
        idx_relativo = idx_faixa - 1  # 0, 1, 2, 3
        if idx_relativo % 2 == 0:
            cor_fundo = "D9D9D9"  # RGB(217,217,217) - Primeira linha
        else:
            cor_fundo = "FFFFFF"  # RGB(255,255,255) - Segunda linha (branco)
        
        # Coluna 1: Descrição
        row.cells[0].text = descricao
        paragraph = row.cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[0]._element.get_or_add_tcPr().append(shading)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Coluna 2: Quantidade (NEGRITO)
        row.cells[1].text = str(quantidade)
        paragraph = row.cells[1].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True  # NEGRITO
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[1]._element.get_or_add_tcPr().append(shading)
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Coluna 3: Percentual (NEGRITO)
        percentual = round((quantidade / total * 100)) if total > 0 else 0
        row.cells[2].text = f"{percentual}%"
        paragraph = row.cells[2].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True  # NEGRITO
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[2]._element.get_or_add_tcPr().append(shading)
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Altura da linha
        table.rows[idx_faixa].height = Cm(1.0)
    
    # === LINHA TOTAL ===
    row_total = table.rows[5]
    
    row_total.cells[0].text = 'Total'
    paragraph = row_total.cells[0].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    shading = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
    row_total.cells[0]._element.get_or_add_tcPr().append(shading)
    row_total.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Coluna quantidade (NEGRITO)
    row_total.cells[1].text = str(total)
    paragraph = row_total.cells[1].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    shading = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
    row_total.cells[1]._element.get_or_add_tcPr().append(shading)
    row_total.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Coluna percentual (NEGRITO)
    row_total.cells[2].text = '100%'
    paragraph = row_total.cells[2].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    shading = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
    row_total.cells[2]._element.get_or_add_tcPr().append(shading)
    row_total.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[5].height = Cm(1.0)
    
    # Ajustar largura das colunas
    for row in table.rows:
        row.cells[0].width = Cm(8.1)    # Coluna descrição
        row.cells[1].width = Cm(2.75)   # Coluna quantidade
        row.cells[2].width = Cm(4.0)    # Coluna percentual
    
    # === APLICAR RECUO DE +1,5cm NA TABELA ===
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Adicionar recuo de +1,5cm (positivo = para direita)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(1.5 * 567)))  # Converter cm para twips (1cm = 567 twips)
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Adicionar espaçamento após a tabela
    doc.add_paragraph()
