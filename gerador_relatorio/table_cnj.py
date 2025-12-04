"""
Módulo para tabelas de Metas Nacionais do CNJ
Contém funções para adicionar tabelas com dados do CNJ incluindo mesclagem de células
"""

import pandas as pd
import re
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from .config import Config


def adicionar_tabela_metas_nacionais(doc):
    """
    Adiciona tabela de metas nacionais do CNJ.
    Mostra META NACIONAL, INSTÂNCIA e PERCENTUAL DE CUMPRIMENTO.
    Mescla células da mesma meta.
    """
    # Carregar dados do CNJ
    try:
        df_cnj = pd.read_excel('exports/resultados_cnj.xlsx')
    except FileNotFoundError:
        print("⚠️  Arquivo 'resultados_cnj.xlsx' não encontrado. Tabela de metas nacionais não será adicionada.")
        return
    
    # Adicionar parágrafo de espaçamento
    doc.add_paragraph()
    
    # Calcular número de linhas (cabeçalho + dados)
    num_linhas_dados = len(df_cnj)
    num_linhas = 1 + num_linhas_dados  # Apenas cabeçalho + dados
    
    # Criar tabela
    table = doc.add_table(rows=num_linhas, cols=3)
    table.style = 'Table Grid'
    
    # === LINHA 1: CABEÇALHO ===
    header_row = table.rows[0]
    headers = ['META NACIONAL', 'INSTÂNCIA', 'PERCENTUAL DE CUMPRIMENTO']
    
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
            run.font.color.rgb = RGBColor(255, 255, 255)  # Texto branco
        
        # Cor de fundo RGB(228,108,10)
        shading = parse_xml(r'<w:shd {} w:fill="E46C0A"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[0].height = Cm(0.8)
    
    # === PREPARAR DADOS PARA MESCLAGEM ===
    # Agrupar por Meta para identificar células a mesclar
    meta_ranges = {}  # {meta: (linha_inicial, linha_final, texto_meta)}
    meta_anterior = None
    linha_inicial_meta = 1  # Agora começa em 1 (linha após cabeçalho)
    
    for idx, (_, row_data) in enumerate(df_cnj.iterrows(), start=1):  # Começa em 1
        meta_completa = row_data.get('Meta', 'N/D')
        
        if meta_completa != meta_anterior:
            if meta_anterior is not None:
                # Salvar range da meta anterior (idx - 1 porque idx já é a nova meta)
                meta_ranges[meta_anterior] = (linha_inicial_meta, idx - 1)
            meta_anterior = meta_completa
            linha_inicial_meta = idx
    
    # Salvar a última meta
    if meta_anterior is not None:
        meta_ranges[meta_anterior] = (linha_inicial_meta, num_linhas_dados)
    
    # === LINHAS DE DADOS ===
    meta_atual = None
    for idx_linha, (_, row_data) in enumerate(df_cnj.iterrows(), start=1):  # Começa em 1
        row = table.rows[idx_linha]
        
        # Extrair dados
        meta_completa = row_data.get('Meta', 'N/D')
        categoria = row_data.get('Categoria', 'Total')
        resultado = row_data.get('Resultado', 'N/D')
        
        # Coluna 1: META NACIONAL (só preenche na primeira linha de cada meta)
        if meta_completa != meta_atual:
            meta_atual = meta_completa
            
            # Extrair apenas o número da meta
            match = re.search(r'Meta\s*(\d+)', meta_completa)
            numero_meta = match.group(1) if match else "?"
            
            # Montar texto: CNJ {numero} - {subtítulo} - {descrição}
            subtitulo = row_data.get('Subtítulo', '')
            descricao = row_data.get('Descrição', '')
            
            # Começar com "CNJ {numero}"
            texto_cnj = f"CNJ {numero_meta}"
            texto_complemento = ""
            
            # Adicionar subtítulo se não for N/D (removendo quebras de linha)
            if pd.notna(subtitulo) and subtitulo != 'N/D':
                subtitulo_limpo = str(subtitulo).replace('\n', ' ').replace('\r', ' ').strip()
                texto_complemento += f" - {subtitulo_limpo}"
            
            # Adicionar descrição se disponível (removendo quebras de linha)
            if pd.notna(descricao) and descricao != 'Descrição não encontrada':
                descricao_limpa = str(descricao).replace('\n', ' ').replace('\r', ' ').strip()
                texto_complemento += f" - {descricao_limpa}"
            
            # Limpar parágrafo e adicionar texto formatado
            row.cells[0].text = ""
            paragraph = row.cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Remover espaçamentos do parágrafo para deixar margens próximas ao texto
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Adicionar "CNJ {numero}" em negrito
            run_cnj = paragraph.add_run(texto_cnj)
            run_cnj.font.size = Pt(11)
            run_cnj.font.bold = True
            run_cnj.font.name = Config.FONTE_PADRAO
            
            # Adicionar resto do texto em fonte normal (sem quebra de linha)
            if texto_complemento:
                run_complemento = paragraph.add_run(texto_complemento)
                run_complemento.font.size = Pt(11)
                run_complemento.font.name = Config.FONTE_PADRAO
            
            # NÃO centralizar verticalmente, deixar no topo
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        else:
            # Limpa completamente o conteúdo das células que serão mescladas
            cell = row.cells[0]
            # Remover todos os parágrafos
            for paragraph in cell.paragraphs[1:]:
                p = paragraph._element
                p.getparent().remove(p)
            # Limpar o primeiro parágrafo
            cell.paragraphs[0].text = ""
        
        # Coluna 2: INSTÂNCIA
        # Remover duplicação (ex: "1º Grau1º Grau" -> "1º Grau")
        categoria_limpa = categoria
        if len(categoria) > 0 and len(categoria) % 2 == 0:
            metade = len(categoria) // 2
            primeira_metade = categoria[:metade]
            segunda_metade = categoria[metade:]
            if primeira_metade == segunda_metade:
                categoria_limpa = primeira_metade
        
        row.cells[1].text = categoria_limpa
        paragraph = row.cells[1].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = Config.FONTE_PADRAO
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Coluna 3: PERCENTUAL DE CUMPRIMENTO
        row.cells[2].text = str(resultado)
        paragraph = row.cells[2].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = Config.FONTE_PADRAO
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # === MESCLAR CÉLULAS DA COLUNA META NACIONAL E AJUSTAR ALTURAS ===
    for meta, (linha_ini, linha_fim) in meta_ranges.items():
        num_linhas_instancia = linha_fim - linha_ini + 1
        
        if linha_fim > linha_ini:  # Só mescla se houver mais de uma linha
            cell_inicial = table.rows[linha_ini].cells[0]
            cell_final = table.rows[linha_fim].cells[0]
            cell_inicial.merge(cell_final)
            
            # Remover parágrafos vazios extras da célula mesclada
            paragraphs_to_remove = []
            for i, paragraph in enumerate(cell_inicial.paragraphs):
                if i > 0:  # Manter apenas o primeiro parágrafo
                    paragraphs_to_remove.append(paragraph)
            
            for paragraph in paragraphs_to_remove:
                p = paragraph._element
                p.getparent().remove(p)
            
            # Centralizar verticalmente o texto na célula mesclada
            cell_inicial.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # === AJUSTAR ALTURA DAS LINHAS ===
        # Calcular altura mínima necessária baseada nas linhas de Instância
        altura_minima_instancia = Cm(num_linhas_instancia * 0.7)
        
        # Definir altura para cada linha do range
        for idx_linha in range(linha_ini, linha_fim + 1):
            # A altura de cada linha individual deve ser proporcional
            table.rows[idx_linha].height = Cm(0.7)
    
    # === DEFINIR LARGURAS DAS COLUNAS ===
    # Forçar larguras específicas das colunas
    table.columns[0].width = Cm(10.5)  # META NACIONAL
    table.columns[1].width = Cm(3.75)  # INSTÂNCIA
    table.columns[2].width = Cm(3.75)  # PERCENTUAL DE CUMPRIMENTO
    
    # Garantir larguras também em cada célula
    for row in table.rows:
        row.cells[0].width = Cm(10.5)  # META NACIONAL
        row.cells[1].width = Cm(3.75)  # INSTÂNCIA
        row.cells[2].width = Cm(3.75)  # PERCENTUAL DE CUMPRIMENTO
    
    # === APLICAR BORDAS PERSONALIZADAS ===
    # Remover todas as bordas e adicionar apenas borda inferior de 1pt
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            
            # Remover bordas existentes se houver
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)
            
            # Criar novas bordas
            tcBorders = OxmlElement('w:tcBorders')
            
            # Borda superior: none
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'none')
            top.set(qn('w:sz'), '0')
            tcBorders.append(top)
            
            # Borda esquerda: none
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'none')
            left.set(qn('w:sz'), '0')
            tcBorders.append(left)
            
            # Borda inferior: single 1pt
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')  # 1pt = 8 eighths of a point
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)
            
            # Borda direita: none
            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'none')
            right.set(qn('w:sz'), '0')
            tcBorders.append(right)
            
            tcPr.append(tcBorders)
    
    # === SEM RECUO PARA ESTA TABELA ===
    # Tabela de metas nacionais não tem recuo aplicado
    
    print(f"✅ Tabela de Metas Nacionais adicionada ({num_linhas_dados} linhas)")
