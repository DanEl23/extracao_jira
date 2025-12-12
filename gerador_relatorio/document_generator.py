"""
Módulo para geração do documento final integrando templates e dados
Responsável por criar o relatório Word combinando estrutura, conteúdo e dados
"""

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, List, Any

from .config import Config
from .document_builder import criar_documento, criar_secao_paisagem_inicial
from .table_historico import adicionar_tabela_historica, adicionar_tabela_macrodesafio
from .table_cnj import adicionar_tabela_metas_nacionais
from .table_monitoramento import adicionar_tabela_resultado_monitoramento
from .table_superintendencia import (
    adicionar_nova_secao_superintendencia,
    adicionar_secao_macrodesafio
)


class DocumentGenerator:
    """Gerador de documento integrando templates e dados"""
    
    def __init__(self, template_data: Dict, df, grupos_super: Dict, variaveis: Dict):
        """
        Inicializa o gerador
        
        Args:
            template_data: Dados processados dos templates (estrutura + conteúdo)
            df: DataFrame com dados das metas
            grupos_super: Dicionário com grupos por superintendência
            variaveis: Dicionário com variáveis calculadas
        """
        self.estrutura = template_data['estrutura']
        self.conteudo = template_data['conteudo']
        self.df = df
        self.grupos_super = grupos_super
        self.variaveis = variaveis
        self.doc = None
        self.secao_paisagem_criada = False
    
    def gerar_documento(self):
        """
        Gera documento final integrando templates e dados
        
        Returns:
            Documento Word gerado
        """
        # Criar documento base (primeira página em retrato)
        print("📝 Criando documento base...")
        primeira_super = list(self.grupos_super.keys())[0] if self.grupos_super else 'Presidência'
        self.doc = criar_documento(primeira_super)
        
        # Remover cabeçalho e rodapé da primeira seção (sumário)
        primeira_secao = self.doc.sections[0]
        
        # Limpar cabeçalho
        header = primeira_secao.header
        for paragraph in list(header.paragraphs):
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Limpar rodapé
        footer = primeira_secao.footer
        for paragraph in list(footer.paragraphs):
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        # Adicionar campo TOC Word (para teste de referência automática de páginas)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        toc_paragraph = self.doc.add_paragraph()
        run = toc_paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        toc_paragraph.paragraph_format.space_after = Pt(12)
        toc_paragraph.paragraph_format.space_before = Pt(0)
        toc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Adicionar sumário
        print("📑 Gerando sumário...")
        self._adicionar_sumario()
        
        # Adicionar sumário detalhado das metas
        print("📋 Gerando sumário detalhado das metas...")
        self._adicionar_sumario_metas()
        
        # Adicionar quebra de página e nova seção para o conteúdo (com cabeçalho/rodapé)
        self._adicionar_nova_secao_conteudo()
        
        # Iterar pela estrutura do sumário
        print("\n✍️  Processando estrutura do template...")
        total = len(self.estrutura)
        
        for idx, item in enumerate(self.estrutura):
            chave = item['chave']
            level = item['level']
            texto = item['texto']
            prefixo = item['prefixo']
            
            print(f"   [{idx+1}/{total}] {chave}")
            
            # Adicionar título formatado
            self._adicionar_titulo(texto, prefixo, level)
            
            # Buscar e processar conteúdo correspondente
            if chave in self.conteudo:
                self._processar_conteudo(chave)
            else:
                # Título sem conteúdo (apenas título fica no doc)
                pass
        
        print("\n✅ Documento gerado com sucesso!")
        return self.doc
    
    def _adicionar_numero_pagina(self, paragrafo):
        """
        Adiciona numeração de páginas no formato 'Página X de Y'
        
        Args:
            paragrafo: Parágrafo onde adicionar a numeração
        """
        run = paragrafo.add_run()
        
        # Adicionar texto "Página "
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # Adicionar " de "
        run.add_text(' de ')
        
        # Adicionar total de páginas
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    def _adicionar_titulo(self, texto: str, prefixo: str, level: int):
        """
        Adiciona título formatado ao documento
        
        Args:
            texto: Texto do título
            prefixo: Prefixo numérico (1, 1.1, etc.)
            level: Nível hierárquico
        """
        # Definir estilo de título do Word conforme o nível
        if level == 1:
            style = 'Heading 1'
        elif level == 2:
            style = 'Heading 2'
        else:
            style = 'Heading 3'

        para = self.doc.add_paragraph(style=style)

        # Formatar título baseado no level (mantém visual)
        if level == 1:
            texto_completo = f"{prefixo}. {texto}"
            run = para.add_run(texto_completo)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(227, 108, 10)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.left_indent = Cm(0.5)
        elif level == 2:
            texto_completo = f"{prefixo} {texto}"
            run = para.add_run(texto_completo)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(227, 108, 10)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.left_indent = Cm(0.75)
        else:
            texto_completo = f"{prefixo} {texto}"
            run = para.add_run(texto_completo)
            run.font.size = Pt(11)
            run.font.bold = True
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)

        run.font.name = Config.FONTE_PADRAO
    
    def _processar_conteudo(self, chave: str):
        """
        Processa todos os blocos de conteúdo de uma seção
        
        Args:
            chave: Chave do título (ex: "1 INTRODUÇÃO")
        """
        blocos = self.conteudo[chave]
        
        for bloco in blocos:
            tipo = bloco['tipo']
            
            if tipo == 'PARAGRAFO':
                self._adicionar_paragrafo(bloco)
            
            elif tipo == 'PARAGRAFO_COM_VARIAVEL':
                self._adicionar_paragrafo_variavel(bloco)
            
            elif tipo == 'TEXTO_DESTAQUE':
                self._adicionar_texto_destaque(bloco)
            
            elif tipo == 'LISTA_NUMERADA':
                self._adicionar_lista_numerada(bloco)
            
            elif tipo == 'LISTA_MARCADORES':
                self._adicionar_lista_marcadores(bloco)
            
            elif tipo == 'TABELA_HISTORICA':
                adicionar_tabela_historica(self.doc)
            
            elif tipo == 'TABELA_MACRODESAFIO':
                adicionar_tabela_macrodesafio(self.doc)
            
            elif tipo == 'TABELA_MONITORAMENTO':
                adicionar_tabela_resultado_monitoramento(self.doc)
            
            elif tipo == 'TABELA_CNJ':
                adicionar_tabela_metas_nacionais(self.doc)
            
            elif tipo == 'SECAO_SUPERINTENDENCIAS':
                self._gerar_secoes_superintendencias()
            
            elif tipo == 'QUEBRA_PAGINA':
                self._adicionar_quebra_pagina()
    
    def _adicionar_paragrafo(self, bloco: Dict):
        """Adiciona parágrafo normal"""
        para = self.doc.add_paragraph()
        self._processar_texto_com_formatacao_inline(para, bloco['texto'])
        
        # Aplicar alinhamento
        alinhamento = bloco.get('alinhamento', 'JUSTIFY')
        if alinhamento == 'CENTER':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alinhamento == 'RIGHT':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alinhamento == 'JUSTIFY':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.left_indent = Cm(0.5)  # Recuo do bloco
        para.paragraph_format.first_line_indent = Cm(1.27)  # Recuo de primeira linha
    
    def _adicionar_paragrafo_variavel(self, bloco: Dict):
        """Adiciona parágrafo substituindo variáveis"""
        texto = bloco['texto']
        
        # Substituir cada variável encontrada
        for marcador in bloco['variaveis']:
            valor = self._obter_valor_variavel(marcador)
            texto = texto.replace(marcador, str(valor))
        
        # Criar parágrafo com texto substituído
        para = self.doc.add_paragraph()
        self._processar_texto_com_formatacao_inline(para, texto)
        
        # Aplicar alinhamento
        alinhamento = bloco.get('alinhamento', 'JUSTIFY')
        if alinhamento == 'CENTER':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alinhamento == 'RIGHT':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alinhamento == 'JUSTIFY':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.left_indent = Cm(0.5)  # Recuo do bloco
        para.paragraph_format.first_line_indent = Cm(1.27)  # Recuo de primeira linha
    
    def _processar_texto_com_formatacao_inline(self, para, texto: str):
        """Processa texto com marcadores inline **texto** para negrito
        
        Args:
            para: Parágrafo do documento
            texto: Texto com possíveis marcadores **texto**
        """
        import re
        
        # Dividir texto em partes: normal e negrito
        # Padrão: captura texto fora de ** e dentro de **
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        
        for parte in partes:
            if not parte:  # Ignorar strings vazias
                continue
                
            if parte.startswith('**') and parte.endswith('**'):
                # Texto em negrito (remover os **)
                texto_negrito = parte[2:-2]
                run = para.add_run(texto_negrito)
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = Config.FONTE_PADRAO
            else:
                # Texto normal
                run = para.add_run(parte)
                run.font.size = Pt(12)
                run.font.name = Config.FONTE_PADRAO
    
    def _adicionar_texto_destaque(self, bloco: Dict):
        """Adiciona texto destacado (apenas negrito)"""
        para = self.doc.add_paragraph()
        run = para.add_run(bloco['texto'])
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = Config.FONTE_PADRAO
        
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.left_indent = Cm(0.5)  # Recuo de 0,5cm
    
    def _adicionar_lista_numerada(self, bloco: Dict):
        """Adiciona lista numerada"""
        for idx, item in enumerate(bloco['itens'], 1):
            para = self.doc.add_paragraph(f"{idx}. {item}")
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.space_after = Pt(3)
            
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = Config.FONTE_PADRAO
    
    def _adicionar_lista_marcadores(self, bloco: Dict):
        """Adiciona lista com marcadores"""
        for item in bloco['itens']:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.paragraph_format.space_after = Pt(3)
            
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = Config.FONTE_PADRAO
    
    def _adicionar_quebra_pagina(self):
        """Adiciona quebra de página sem criar linha vazia"""
        from docx.enum.text import WD_BREAK
        from docx.shared import Pt
        
        # Criar parágrafo com espaçamento zero
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        # Adicionar quebra de página
        p.add_run().add_break(WD_BREAK.PAGE)
    
    def _gerar_secoes_superintendencias(self):
        """Gera seções detalhadas por superintendência (tabelas dinâmicas em paisagem)"""
        print("\n📋 Gerando seções de superintendências...")
        
        # Criar seção paisagem se ainda não foi criada
        if not self.secao_paisagem_criada:
            primeira_super = list(self.grupos_super.keys())[0] if self.grupos_super else 'Presidência'
            criar_secao_paisagem_inicial(self.doc, primeira_super)
            self.secao_paisagem_criada = True
        
        primeira_super = True
        
        for superintendencia, grupos_macro in self.grupos_super.items():
            print(f"   → {superintendencia}")
            if not primeira_super:
                adicionar_nova_secao_superintendencia(self.doc, superintendencia, False)
            primeira_secao = True
            for macrodesafio, df_grupo in grupos_macro:
                meta_bookmarks = getattr(self, '_meta_bookmarks', None)
                adicionar_secao_macrodesafio(
                    self.doc,
                    macrodesafio,
                    df_grupo,
                    primeira_secao,
                    meta_bookmarks=meta_bookmarks,
                    superintendencia=superintendencia
                )
                primeira_secao = False
            primeira_super = False
    
    def _obter_valor_variavel(self, marcador: str) -> Any:
        """
        Obtém valor de uma variável pelo marcador
        
        Args:
            marcador: Marcador da variável (ex: [NUMERO_METAS_2025])
            
        Returns:
            Valor da variável ou '???' se não encontrada
        """
        from .content_mapper import ContentMapper
        
        # Mapear marcador para nome de variável
        nome_var = ContentMapper.MARCADORES_VARIAVEIS.get(marcador)
        
        if nome_var and nome_var in self.variaveis:
            valor = self.variaveis[nome_var]
            
            # Formatar decimais
            if isinstance(valor, float):
                return f"{valor:.1f}"
            
            return valor
        
        return '???'
    
    def _adicionar_nova_secao_conteudo(self):
        """Adiciona nova seção para conteúdo com cabeçalho e rodapé"""
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Criar nova seção em retrato
        nova_secao = self.doc.add_section(WD_SECTION.NEW_PAGE)
        
        # Copiar configurações da primeira seção
        primeira_secao = self.doc.sections[0]
        nova_secao.page_width = primeira_secao.page_width
        nova_secao.page_height = primeira_secao.page_height
        nova_secao.orientation = primeira_secao.orientation
        nova_secao.top_margin = primeira_secao.top_margin
        nova_secao.bottom_margin = primeira_secao.bottom_margin
        nova_secao.left_margin = primeira_secao.left_margin
        nova_secao.right_margin = primeira_secao.right_margin
        nova_secao.header_distance = primeira_secao.header_distance
        nova_secao.footer_distance = primeira_secao.footer_distance
        
        # Desvincular cabeçalho e rodapé
        nova_secao.header.is_linked_to_previous = False
        nova_secao.footer.is_linked_to_previous = False
        
        # Adicionar cabeçalho
        header = nova_secao.header
        
        # Linha 1: MONITORAMENTO DE METAS ESTRATÉGICAS - 2024
        p1 = header.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run('MONITORAMENTO DE METAS ESTRATÉGICAS - 2024')
        run1.font.size = Pt(11)
        run1.font.bold = True
        run1.font.name = Config.FONTE_PADRAO
        p1.paragraph_format.space_after = Pt(0)
        
        # Linha 2: Relatório Técnico
        p2 = header.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run('Relatório Técnico ao Comitê de Governança e Gestão Estratégica')
        run2.font.size = Pt(11)
        run2.font.name = Config.FONTE_PADRAO
        p2.paragraph_format.space_after = Pt(6)
        
        # Adicionar rodapé
        footer = nova_secao.footer
        
        # Criar tabela de 1 linha x 2 colunas para rodapé compacto
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        
        table = footer.add_table(rows=1, cols=2, width=Cm(19))
        table.autofit = False
        table.allow_autofit = False
        
        # Aplicar recuo negativo para estender além das margens
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '-850')  # -1.5cm em twips (1cm = 567 twips)
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        # Forçar layout fixo
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Configurar larguras das colunas em twips (1cm = 567 twips)
        table.columns[0].width = Cm(18)  # ASPLAG e DEPLAG = 10206 twips
        table.columns[1].width = Cm(1)   # Número página = 567 twips
        
        # Configurar células
        row = table.rows[0]
        
        # Célula 1: ASPLAG e DEPLAG (esquerda)
        cell1 = row.cells[0]
        cell1.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Definir largura explícita da célula 1: 18cm = 10206 twips
        tcPr1 = cell1._element.get_or_add_tcPr()
        tcW1 = OxmlElement('w:tcW')
        tcW1.set(qn('w:w'), '10206')  # 18cm em twips
        tcW1.set(qn('w:type'), 'dxa')
        tcPr1.append(tcW1)
        # Desabilitar quebra automática
        noWrap = OxmlElement('w:noWrap')
        tcPr1.append(noWrap)
        
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p1.add_run('Assessoria Técnica e Jurídica ao Planejamento e à Gestão Institucional - ASPLAG')
        run1.font.size = Pt(9)
        run1.font.name = Config.FONTE_PADRAO
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        # Segunda linha - DEPLAG
        p1_linha2 = cell1.add_paragraph()
        p1_linha2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1_linha2 = p1_linha2.add_run('Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional - DEPLAG')
        run1_linha2.font.size = Pt(9)
        run1_linha2.font.name = Config.FONTE_PADRAO
        p1_linha2.paragraph_format.space_before = Pt(0)
        p1_linha2.paragraph_format.space_after = Pt(0)
        
        # Célula 2: Número da página (direita)
        cell2 = row.cells[1]
        cell2.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Definir largura explícita da célula 2: 1cm = 567 twips
        tcPr2 = cell2._element.get_or_add_tcPr()
        tcW2 = OxmlElement('w:tcW')
        tcW2.set(qn('w:w'), '567')  # 1cm em twips
        tcW2.set(qn('w:type'), 'dxa')
        tcPr2.append(tcW2)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = p2.add_run()
        run2.font.size = Pt(12)  # Aumentado de 9pt para 12pt
        run2.font.name = Config.FONTE_PADRAO
        # Adicionar campo PAGE
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        
        # Remover bordas laterais e inferior, manter apenas borda superior
        for i, cell in enumerate(row.cells):
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            # Borda superior (linha fina)
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')  # 0.5pt
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), '000000')
            tcBorders.append(top)
            
            # Sem bordas laterais e inferior
            for border in ['left', 'bottom', 'right']:
                elem = OxmlElement(f'w:{border}')
                elem.set(qn('w:val'), 'none')
                tcBorders.append(elem)
                
            tcPr.append(tcBorders)
            
            # Remover padding interno para altura mínima
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'bottom', 'left', 'right']:
                mar = OxmlElement(f'w:{margin}')
                mar.set(qn('w:w'), '0')
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            tcPr.append(tcMar)
    
    def _adicionar_sumario(self):
        """Adiciona sumário ao documento baseado na estrutura extraída"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Título "SUMÁRIO"
        titulo_sumario = self.doc.add_paragraph()
        titulo_sumario.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = titulo_sumario.add_run('SUMÁRIO')
        run_titulo.font.size = Pt(16)
        run_titulo.font.bold = True
        run_titulo.font.name = Config.FONTE_PADRAO
        titulo_sumario.paragraph_format.space_before = Pt(12)
        titulo_sumario.paragraph_format.space_after = Pt(0)
        
        # Adicionar itens do sumário
        for item in self.estrutura:
            prefixo = item['prefixo']
            texto = item['texto']
            level = item['level']
            
            # Criar linha do sumário
            para = self.doc.add_paragraph()
            
            # Configurar recuo baseado no nível
            if level == 1:
                para.paragraph_format.left_indent = Cm(0)
            elif level == 2:
                para.paragraph_format.left_indent = Cm(0.2)
            else:
                para.paragraph_format.left_indent = Cm(1.0)
            
            # Adicionar texto do sumário (prefixo + título)
            texto_sumario = f"{prefixo}   {texto}"
            run_texto = para.add_run(texto_sumario)
            run_texto.font.size = Pt(10)
            run_texto.font.name = Config.FONTE_PADRAO
            
            # Negrito para níveis 1 e 2
            if level in [1, 2]:
                run_texto.font.bold = True
            
            # Espaçamento baseado no nível
            if level == 1:
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
            elif level == 2:
                para.paragraph_format.space_after = Pt(2)
            
            # Adicionar TAB com pontos de preenchimento
            # Criar elemento de tabulação
            pPr = para._element.get_or_add_pPr()
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = OxmlElement('w:tabs')
                pPr.append(tabs)
            
            # Adicionar tab stop próximo da margem direita
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '9072')  # 16cm exatos (16 * 567 = 9072 twips)
            tab.set(qn('w:leader'), 'dot')  # Pontos de preenchimento
            tabs.append(tab)
            
            # Adicionar TAB e número de página
            run_tab = para.add_run('\t')
            run_numero = para.add_run('1')  # Placeholder - Word atualizará
            run_numero.font.size = Pt(10)
            run_numero.font.name = Config.FONTE_PADRAO
            
            para.paragraph_format.space_after = Pt(0)
    
    def _adicionar_sumario_metas(self):
        """Adiciona sumário detalhado das metas por superintendência"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import pandas as pd
        from base_dados import ORDEM_SUPERINTENDENCIAS
        
        # Espaçamento antes do sumário de metas
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # Iterar por cada superintendência na ordem correta
        for super_nome in ORDEM_SUPERINTENDENCIAS:
            lista_grupos = self.grupos_super.get(super_nome, [])
            if not lista_grupos:
                continue
            dfs = [df_macro for _, df_macro in lista_grupos]
            df_super = pd.concat(dfs, ignore_index=True) if dfs else pd.DataFrame()
            if df_super.empty:
                continue
            para_super = self.doc.add_paragraph()
            run_super = para_super.add_run(super_nome.upper())
            run_super.font.size = Pt(10)
            run_super.font.bold = True
            run_super.font.name = Config.FONTE_PADRAO
            para_super.paragraph_format.space_before = Pt(6)
            para_super.paragraph_format.space_after = Pt(3)

            for global_idx, (idx, row) in enumerate(df_super.iterrows()):
                texto_meta = row.get(Config.COLUNAS['METAKEY'], 'N/A')
                # Gerar um nome de bookmark único e válido para cada meta
                safe_super = ''.join(c if c.isalnum() else '_' for c in str(super_nome))
                bookmark_name = f"meta_{safe_super}_{global_idx}"

                # Criar linha da meta no sumário
                para_meta = self.doc.add_paragraph()
                para_meta.paragraph_format.left_indent = Cm(0.2)
                run_meta = para_meta.add_run(texto_meta)
                run_meta.font.size = Pt(10)
                run_meta.font.name = Config.FONTE_PADRAO

                # Tabulação com pontos
                pPr = para_meta._element.get_or_add_pPr()
                tabs = pPr.find(qn('w:tabs'))
                if tabs is None:
                    tabs = OxmlElement('w:tabs')
                    pPr.append(tabs)
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9072')
                tab.set(qn('w:leader'), 'dot')
                tabs.append(tab)

                para_meta.add_run('\t')
                # Adicionar campo PAGEREF para o bookmark
                run_pageref = para_meta.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = f'PAGEREF {bookmark_name} \\h'
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                run_pageref._r.append(fldChar1)
                run_pageref._r.append(instrText)
                run_pageref._r.append(fldChar2)
                run_pageref._r.append(fldChar3)
                run_pageref.font.size = Pt(10)
                run_pageref.font.name = Config.FONTE_PADRAO
                para_meta.paragraph_format.space_after = Pt(0)

                # Salvar o bookmark_name para uso posterior
                if not hasattr(self, '_meta_bookmarks'):
                    self._meta_bookmarks = {}
                self._meta_bookmarks.setdefault(super_nome, {})[texto_meta] = bookmark_name
