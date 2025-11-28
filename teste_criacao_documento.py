import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- 1. Carregar os dados ---
# Substitua pelo caminho correto do seu arquivo se necessário
file_path = 'teste_integração.xlsx' 
try:
    df = pd.read_excel(file_path)
except:
    # Se o arquivo original não for encontrado, tenta ler o CSV
    df = pd.read_csv('teste_integração.xlsx - Sheet1.csv')

df = df.fillna('') # Remove valores nulos

# --- 2. Preparar o Documento ---
doc = Document()

# Configuração de Página: PAISAGEM (Landscape)
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
# Inverte as dimensões padrão para refletir a orientação paisagem (A4)
section.page_width = Inches(11.69)
section.page_height = Inches(8.27)
# Margens estreitas para aproveitar melhor o espaço
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

# --- 3. Funções Auxiliares de Estilo ---

def set_cell_background(cell, color_hex):
    """Aplica cor de fundo (shading) à célula via XML."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_cell_content(cell, label, value, label_bold=True):
    """Adiciona rótulo e valor na célula com formatação."""
    # Limpa parágrafos vazios existentes
    cell.text = "" 
    p = cell.add_paragraph()
    
    # Rótulo em Negrito e Caixa Alta
    run_label = p.add_run(f"{label.upper()}: ")
    run_label.bold = label_bold
    run_label.font.name = 'Arial'
    run_label.font.size = Pt(10)
    
    # Valor normal
    run_value = p.add_run(str(value))
    run_value.font.name = 'Arial'
    run_value.font.size = Pt(10)
    
    # Espaçamento para não ficar colado nas bordas
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def set_macro_header(cell, text):
    """Estilo específico para o cabeçalho Macrodesafio."""
    cell.text = ""
    p = cell.add_paragraph()
    run = p.add_run(f"MACRODESAFIO: {text}")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0) # Preto
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# --- 4. Gerar as Tabelas ---
for index, row in df.iterrows():
    
    # Cria tabela 5 linhas x 2 colunas
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid' # Garante as bordas padrão do Word
    table.autofit = False 
    
    # Ajuste largura das colunas (metade da página para cada)
    # Largura útil ~10.69 inches (11.69 - 1.0 margem)
    col_width = Inches(5.3)
    for row_table in table.rows:
        for cell in row_table.cells:
            cell.width = col_width

    # --- LINHA 1: Macrodesafio (Mesclada) ---
    cell_macro = table.cell(0, 0)
    cell_macro.merge(table.cell(0, 1))
    set_cell_background(cell_macro, "D9D9D9") # Cinza médio
    set_macro_header(cell_macro, row['Macrodesafio'])

    # --- LINHA 2: Indicador e Meta ---
    # Indicador
    format_cell_content(table.cell(1, 0), "Indicador", row['Indicador'])
    # Meta
    format_cell_content(table.cell(1, 1), "Meta", row['MetaKey'])

    # --- LINHA 3: Unidade e Polaridade ---
    # Unidade
    format_cell_content(table.cell(2, 0), "Unidade Responsável", row['Unidade Gestora'])
    # Polaridade
    format_cell_content(table.cell(2, 1), "Polaridade", row['Polaridade'])

    # --- LINHA 4: Resultado e Iniciativa ---
    # Resultado
    format_cell_content(table.cell(3, 0), "Resultado Apurado", row['Valor Apurado'])
    # Iniciativa
    format_cell_content(table.cell(3, 1), "Iniciativa(s)", row['Iniciativa'])

    # --- LINHA 5: Situação (Mesclada) ---
    cell_sit = table.cell(4, 0)
    cell_sit.merge(table.cell(4, 1))
    
    # Situação tem um estilo levemente diferente (texto justificado)
    cell_sit.text = ""
    p = cell_sit.add_paragraph()
    run_lbl = p.add_run("SITUAÇÃO: ")
    run_lbl.bold = True
    run_lbl.font.name = 'Arial'
    run_lbl.font.size = Pt(10)
    
    # Adiciona o texto da situação
    run_val = p.add_run(str(row['Informação complementar texto']))
    run_val.font.name = 'Arial'
    run_val.font.size = Pt(10)
    
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # Adiciona espaço entre as tabelas
    doc.add_paragraph("\n")
    # doc.add_page_break() # Descomente se quiser forçar uma nova página por registro

# --- 5. Salvar ---
output_filename = 'Relatorio_Formatado_Final.docx'
doc.save(output_filename)
print(f"Documento gerado com sucesso: {output_filename}")