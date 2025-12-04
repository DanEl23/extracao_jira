"""
Script para criar os templates de exemplo (Sumario_Modelo.docx e Conteudo_Fonte.docx)
Execute este script para gerar os arquivos .docx de template
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def criar_sumario_modelo():
    """Cria o template Sumario_Modelo.docx"""
    
    print("📄 Criando Sumario_Modelo.docx...")
    
    doc = Document()
    
    # Título do Sumário
    titulo = doc.add_paragraph()
    run_titulo = titulo.add_run('SUMÁRIO')
    run_titulo.font.size = Pt(14)
    run_titulo.font.bold = True
    run_titulo.font.name = 'Arial'
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_after = Pt(18)
    
    # Estrutura hierárquica do relatório
    estrutura = [
        ("1", "INTRODUÇÃO", 1),
        ("1.1", "Contextualização", 2),
        ("1.2", "Objetivos do Relatório", 2),
        ("1.3", "Período de Referência", 2),
        
        ("2", "METAS APROVADAS", 1),
        ("2.1", "Total de Metas Estratégicas", 2),
        ("2.2", "Distribuição por Macrodesafio", 2),
        ("2.3", "Evolução Histórica", 2),
        
        ("3", "RESULTADO DO MONITORAMENTO", 1),
        ("3.1", "Indicadores de Cumprimento", 2),
        ("3.2", "Análise por Faixa de Desempenho", 2),
        ("3.3", "Destaques do Período", 2),
        
        ("4", "METAS NACIONAIS DO CNJ", 1),
        ("4.1", "Visão Geral", 2),
        ("4.2", "Detalhamento por Meta", 2),
        ("4.3", "Alinhamento Institucional", 2),
        
        ("5", "SUPERINTENDÊNCIAS", 1),
        ("5.1", "Presidência", 2),
        ("5.2", "1ª Vice-Presidência", 2),
        ("5.3", "2ª Vice-Presidência", 2),
        ("5.4", "3ª Vice-Presidência", 2),
        ("5.5", "Corregedoria", 2),
        ("5.6", "Escola Judicial", 2),
        
        ("6", "CONSIDERAÇÕES FINAIS", 1),
        ("6.1", "Síntese dos Resultados", 2),
        ("6.2", "Oportunidades de Melhoria", 2),
        ("6.3", "Próximos Passos", 2),
    ]
    
    for prefixo, texto, level in estrutura:
        para = doc.add_paragraph()
        
        # Adicionar recuo baseado no level
        if level == 2:
            para.paragraph_format.left_indent = Cm(1.0)
        elif level == 3:
            para.paragraph_format.left_indent = Cm(2.0)
        
        # Formatação do texto
        run = para.add_run(f"{prefixo} {texto}")
        run.font.name = 'Arial'
        
        if level == 1:
            run.font.size = Pt(12)
            run.font.bold = True
        else:
            run.font.size = Pt(11)
        
        para.paragraph_format.space_after = Pt(6)
    
    # Salvar
    caminho = os.path.join(os.path.dirname(__file__), 'Sumario_Modelo.docx')
    doc.save(caminho)
    print(f"   ✅ Salvo em: {caminho}")
    
    return caminho


def criar_conteudo_fonte():
    """Cria o template Conteudo_Fonte.docx com textos e marcadores"""
    
    print("\n📝 Criando Conteudo_Fonte.docx...")
    
    doc = Document()
    
    # Função auxiliar para adicionar títulos
    def adicionar_titulo(texto, level=1):
        para = doc.add_paragraph()
        run = para.add_run(texto)
        run.font.name = 'Arial'
        run.font.bold = True
        
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(227, 108, 10)  # Laranja
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(12)
        elif level == 2:
            run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(8)
        else:
            run.font.size = Pt(11)
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(6)
    
    # Função auxiliar para adicionar parágrafos
    def adicionar_paragrafo(texto):
        para = doc.add_paragraph(texto)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.name = 'Arial'
    
    # Função para adicionar marcador
    def adicionar_marcador(marcador):
        para = doc.add_paragraph(marcador)
        run = para.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Azul para destacar
        para.paragraph_format.space_after = Pt(3)
    
    # ============================================
    # CONTEÚDO DO TEMPLATE
    # ============================================
    
    # 1 INTRODUÇÃO
    adicionar_titulo("1 INTRODUÇÃO", 1)
    
    adicionar_paragrafo(
        "Este relatório apresenta os resultados do monitoramento das metas estratégicas "
        "do Tribunal de Justiça do Estado de Minas Gerais (TJMG) para o ano de [VAR_ANO_ATUAL], "
        "conforme deliberações do Comitê de Governança e Gestão Estratégica."
    )
    
    adicionar_paragrafo(
        "O documento consolida informações sobre o desempenho institucional, identificando "
        "avanços, desafios e oportunidades de melhoria na execução do planejamento estratégico."
    )
    
    # 1.1 Contextualização
    adicionar_titulo("1.1 Contextualização", 2)
    
    adicionar_paragrafo(
        "O planejamento estratégico do TJMG está alinhado às diretrizes do Conselho Nacional "
        "de Justiça (CNJ) e contempla tanto as metas nacionais quanto as metas institucionais "
        "específicas, estruturadas em macrodesafios que orientam as ações das diversas "
        "unidades administrativas e judiciárias."
    )
    
    adicionar_paragrafo(
        "O processo de monitoramento é realizado periodicamente e tem como objetivo acompanhar "
        "o desempenho institucional, fornecendo subsídios para a tomada de decisão pelos gestores "
        "e para eventuais ajustes no planejamento."
    )
    
    # 1.2 Objetivos do Relatório
    adicionar_titulo("1.2 Objetivos do Relatório", 2)
    
    adicionar_paragrafo("#Os principais objetivos deste documento são:")
    
    adicionar_marcador("[INICIAR_LISTA_MARCADORES]")
    adicionar_paragrafo("Apresentar o panorama geral das metas estratégicas aprovadas")
    adicionar_paragrafo("Analisar os resultados apurados no período de referência")
    adicionar_paragrafo("Identificar metas com desempenho satisfatório e aquelas que demandam atenção")
    adicionar_paragrafo("Fornecer informações consolidadas por superintendência e macrodesafio")
    adicionar_paragrafo("Subsidiar decisões estratégicas e ajustes no planejamento institucional")
    adicionar_marcador("[FINALIZAR_LISTA_MARCADORES]")
    
    # 1.3 Período de Referência
    adicionar_titulo("1.3 Período de Referência", 2)
    
    adicionar_paragrafo(
        "Os dados apresentados referem-se ao monitoramento realizado no ano de [VAR_ANO_ATUAL], "
        "contemplando a apuração dos indicadores estabelecidos para as [NUMERO_METAS_2025] metas "
        "estratégicas aprovadas."
    )
    
    # 2 METAS APROVADAS
    doc.add_page_break()
    adicionar_titulo("2 METAS APROVADAS", 1)
    
    adicionar_paragrafo(
        "Esta seção apresenta o quadro geral das metas estratégicas aprovadas pelo TJMG, "
        "incluindo a evolução histórica e a distribuição por macrodesafios institucionais."
    )
    
    # 2.1 Total de Metas Estratégicas
    adicionar_titulo("2.1 Total de Metas Estratégicas", 2)
    
    adicionar_marcador("[INSERIR_TABELA_HISTORICA]")
    
    adicionar_paragrafo(
        "O quadro acima demonstra a evolução histórica das metas estratégicas aprovadas pelo TJMG "
        "nos últimos 4 anos. Observa-se que para o ano de [VAR_ANO_ATUAL] foram aprovadas "
        "[NUMERO_METAS_2025] metas estratégicas, sendo [NUMERO_METAS_CNJ] metas nacionais do CNJ "
        "e [NUMERO_METAS_TJMG] metas institucionais do TJMG."
    )
    
    # 2.2 Distribuição por Macrodesafio
    adicionar_titulo("2.2 Distribuição por Macrodesafio", 2)
    
    adicionar_marcador("[INSERIR_TABELA_MACRODESAFIO]")
    
    adicionar_paragrafo(
        "A distribuição das metas por macrodesafio reflete as prioridades estratégicas estabelecidas "
        "pela instituição, contemplando [NUMERO_MACRODESAFIOS] eixos temáticos que orientam as ações "
        "institucionais."
    )
    
    # 2.3 Evolução Histórica
    adicionar_titulo("2.3 Evolução Histórica", 2)
    
    adicionar_paragrafo(
        "A análise da evolução histórica demonstra o crescimento e amadurecimento do processo de "
        "planejamento estratégico do TJMG, com ampliação gradual do escopo de metas monitoradas "
        "e aperfeiçoamento dos mecanismos de acompanhamento."
    )
    
    # 3 RESULTADO DO MONITORAMENTO
    doc.add_page_break()
    adicionar_titulo("3 RESULTADO DO MONITORAMENTO", 1)
    
    adicionar_paragrafo(
        "A seguir são apresentados os resultados consolidados do monitoramento realizado no período "
        "de referência, com análise dos indicadores de cumprimento e distribuição por faixas de desempenho."
    )
    
    # 3.1 Indicadores de Cumprimento
    adicionar_titulo("3.1 Indicadores de Cumprimento", 2)
    
    adicionar_marcador("[INSERIR_TABELA_MONITORAMENTO]")
    
    adicionar_paragrafo(
        "A tabela acima consolida os resultados do monitoramento, classificando as metas em faixas "
        "de desempenho de acordo com os percentuais de cumprimento apurados."
    )
    
    # 3.2 Análise por Faixa de Desempenho
    adicionar_titulo("3.2 Análise por Faixa de Desempenho", 2)
    
    adicionar_paragrafo(
        "#Destaque Positivo: [PERCENTUAL_VERDE]% das metas atingiram resultado satisfatório, "
        "situando-se na faixa verde (cumprimento igual ou superior a 90%)."
    )
    
    adicionar_paragrafo(
        "As metas classificadas nas faixas amarela e vermelha demandam atenção especial e podem "
        "requerer ajustes nas estratégias de implementação ou revisão dos indicadores estabelecidos."
    )
    
    # 3.3 Destaques do Período
    adicionar_titulo("3.3 Destaques do Período", 2)
    
    adicionar_paragrafo(
        "Entre os principais destaques do período, identificam-se iniciativas que contribuíram "
        "significativamente para o alcance dos resultados, incluindo:"
    )
    
    adicionar_marcador("[INICIAR_LISTA_NUMERADA]")
    adicionar_paragrafo("Implementação de novas ferramentas tecnológicas de apoio à gestão")
    adicionar_paragrafo("Capacitação de equipes para execução das ações estratégicas")
    adicionar_paragrafo("Aprimoramento dos processos de monitoramento e controle")
    adicionar_paragrafo("Integração entre as diversas superintendências e unidades")
    adicionar_marcador("[FINALIZAR_LISTA_NUMERADA]")
    
    # 4 METAS NACIONAIS DO CNJ
    doc.add_page_break()
    adicionar_titulo("4 METAS NACIONAIS DO CNJ", 1)
    
    adicionar_paragrafo(
        "As metas nacionais estabelecidas pelo Conselho Nacional de Justiça (CNJ) representam "
        "compromissos do Poder Judiciário brasileiro com a sociedade, focando em temas prioritários "
        "como celeridade processual, efetividade da execução e acesso à justiça."
    )
    
    # 4.1 Visão Geral
    adicionar_titulo("4.1 Visão Geral", 2)
    
    adicionar_marcador("[INSERIR_TABELA_CNJ]")
    
    adicionar_paragrafo(
        "O TJMG acompanha sistematicamente o cumprimento das [NUMERO_METAS_CNJ] metas nacionais "
        "estabelecidas pelo CNJ, buscando o alinhamento entre as estratégias institucionais e "
        "as diretrizes nacionais."
    )
    
    # 4.2 Detalhamento por Meta
    adicionar_titulo("4.2 Detalhamento por Meta", 2)
    
    adicionar_paragrafo(
        "Cada meta nacional do CNJ possui indicadores específicos e alvos estabelecidos, sendo "
        "monitorada mensalmente para garantir o acompanhamento tempestivo dos resultados e a "
        "possibilidade de intervenções corretivas quando necessário."
    )
    
    # 4.3 Alinhamento Institucional
    adicionar_titulo("4.3 Alinhamento Institucional", 2)
    
    adicionar_paragrafo(
        "O TJMG promove o alinhamento entre as metas nacionais do CNJ e as metas institucionais, "
        "garantindo coerência estratégica e otimização dos recursos disponíveis para alcance dos "
        "objetivos estabelecidos."
    )
    
    # 5 SUPERINTENDÊNCIAS
    doc.add_page_break()
    adicionar_titulo("5 SUPERINTENDÊNCIAS", 1)
    
    adicionar_paragrafo(
        "A partir desta seção são apresentados os resultados detalhados por superintendência e "
        "macrodesafio, permitindo uma análise mais aprofundada do desempenho de cada área responsável."
    )
    
    adicionar_paragrafo(
        "As tabelas a seguir apresentam informações pormenorizadas sobre cada indicador, incluindo "
        "resultados apurados, polaridade, unidades responsáveis e iniciativas relacionadas."
    )
    
    adicionar_marcador("[INICIAR_SECAO_SUPERINTENDENCIAS]")
    
    adicionar_paragrafo(
        "Nota: As seções detalhadas de cada superintendência são geradas automaticamente a partir "
        "dos dados extraídos, incluindo tabelas em formato paisagem com todas as informações dos "
        "indicadores monitorados."
    )
    
    # 6 CONSIDERAÇÕES FINAIS
    doc.add_page_break()
    adicionar_titulo("6 CONSIDERAÇÕES FINAIS", 1)
    
    adicionar_paragrafo(
        "Este relatório consolida os resultados do monitoramento estratégico do TJMG, demonstrando "
        "o comprometimento institucional com a gestão por resultados e a melhoria contínua dos "
        "serviços prestados à sociedade."
    )
    
    # 6.1 Síntese dos Resultados
    adicionar_titulo("6.1 Síntese dos Resultados", 2)
    
    adicionar_paragrafo(
        "O período analisado demonstra avanços significativos no cumprimento das metas estratégicas, "
        "com [PERCENTUAL_VERDE]% das metas atingindo resultados satisfatórios. Este desempenho "
        "reflete o esforço conjunto de todas as unidades e o engajamento dos colaboradores na "
        "execução do planejamento estratégico."
    )
    
    # 6.2 Oportunidades de Melhoria
    adicionar_titulo("6.2 Oportunidades de Melhoria", 2)
    
    adicionar_paragrafo(
        "Identificam-se oportunidades de melhoria relacionadas às metas que não alcançaram os "
        "resultados esperados, demandando análise criteriosa das causas e implementação de ações "
        "corretivas específicas."
    )
    
    adicionar_paragrafo(
        "O aprimoramento contínuo dos processos de planejamento, execução e monitoramento constitui "
        "elemento fundamental para o alcance dos objetivos institucionais."
    )
    
    # 6.3 Próximos Passos
    adicionar_titulo("6.3 Próximos Passos", 2)
    
    adicionar_paragrafo("#As principais ações previstas para o próximo período incluem:")
    
    adicionar_marcador("[INICIAR_LISTA_MARCADORES]")
    adicionar_paragrafo("Revisão das estratégias para metas com desempenho insatisfatório")
    adicionar_paragrafo("Intensificação do acompanhamento mensal dos indicadores críticos")
    adicionar_paragrafo("Capacitação continuada das equipes responsáveis")
    adicionar_paragrafo("Fortalecimento da governança e comunicação estratégica")
    adicionar_paragrafo("Avaliação de viabilidade de ajustes no planejamento quando necessário")
    adicionar_marcador("[FINALIZAR_LISTA_MARCADORES]")
    
    # Salvar
    caminho = os.path.join(os.path.dirname(__file__), 'Conteudo_Fonte.docx')
    doc.save(caminho)
    print(f"   ✅ Salvo em: {caminho}")
    
    return caminho


def criar_readme_templates():
    """Cria arquivo README explicativo dos templates"""
    
    print("\n📚 Criando README_TEMPLATES.md...")
    
    readme_content = """# 📘 GUIA DE USO DOS TEMPLATES

## 🎯 Visão Geral

Os templates permitem controlar a estrutura e conteúdo do relatório **sem modificar código**.

### **Arquivos:**

1. **Sumario_Modelo.docx** - Define estrutura hierárquica do relatório
2. **Conteudo_Fonte.docx** - Contém textos, marcadores e variáveis

---

## 📝 SUMARIO_MODELO.DOCX

### **Formato:**

```
1 TÍTULO NÍVEL 1
1.1 Título Nível 2
1.2 Outro Título Nível 2
2 OUTRO TÍTULO NÍVEL 1
2.1 Subtítulo
2.1.1 Nível 3
```

### **Regras:**

✅ Use numeração hierárquica (1, 1.1, 1.1.1)  
✅ Ponto final no prefixo é opcional (1. ou 1)  
✅ Títulos devem ser EXATOS - serão mapeados no conteúdo  
❌ Não adicione números de página (gerados automaticamente)  
❌ Não use tabulações ou pontos de preenchimento  

### **Exemplo:**

```
1 INTRODUÇÃO
1.1 Contextualização
1.2 Objetivos do Relatório

2 METAS APROVADAS
2.1 Total de Metas Estratégicas
```

---

## 📄 CONTEUDO_FONTE.DOCX

### **Estrutura:**

O conteúdo deve seguir a mesma hierarquia do sumário. Cada título do sumário deve ter uma seção correspondente no conteúdo.

### **Marcadores de Tabelas Dinâmicas:**

Estes marcadores são substituídos por tabelas geradas automaticamente a partir dos dados:

| Marcador | Descrição |
|----------|-----------|
| `[INSERIR_TABELA_HISTORICA]` | Tabela de metas aprovadas (4 anos) |
| `[INSERIR_TABELA_MACRODESAFIO]` | Distribuição por macrodesafio |
| `[INSERIR_TABELA_MONITORAMENTO]` | Resultado do monitoramento (faixas) |
| `[INSERIR_TABELA_CNJ]` | Metas nacionais do CNJ |
| `[INICIAR_SECAO_SUPERINTENDENCIAS]` | Seções detalhadas por superintendência |

**Uso:**

```
2.1 Total de Metas Estratégicas

[INSERIR_TABELA_HISTORICA]

O quadro acima demonstra a evolução...
```

### **Marcadores de Variáveis:**

Estes marcadores são substituídos por valores calculados automaticamente:

| Marcador | Descrição | Exemplo |
|----------|-----------|---------|
| `[NUMERO_METAS_2025]` | Total de metas do ano atual | 57 |
| `[NUMERO_METAS_CNJ]` | Total de metas nacionais CNJ | 9 |
| `[NUMERO_METAS_TJMG]` | Total de metas institucionais TJMG | 48 |
| `[NUMERO_MACRODESAFIOS]` | Total de macrodesafios | 6 |
| `[PERCENTUAL_VERDE]` | % de metas na faixa verde | 75.5 |
| `[VAR_ANO_ATUAL]` | Ano atual do relatório | 2025 |

**Uso:**

```
Foram aprovadas [NUMERO_METAS_2025] metas para o ano de [VAR_ANO_ATUAL].
```

### **Formatação Especial:**

#### **1. Texto com Destaque (Negrito + Vermelho)**

Inicie a linha com `#` para aplicar formatação de destaque:

```
#Este texto será exibido em negrito e cor vermelha
```

#### **2. Listas Numeradas**

Use marcadores de controle:

```
Os principais objetivos são:

[INICIAR_LISTA_NUMERADA]
Apresentar panorama geral das metas
Analisar resultados apurados
Identificar oportunidades de melhoria
[FINALIZAR_LISTA_NUMERADA]
```

Resultado:
```
1. Apresentar panorama geral das metas
2. Analisar resultados apurados
3. Identificar oportunidades de melhoria
```

#### **3. Listas com Marcadores (Bullets)**

```
[INICIAR_LISTA_MARCADORES]
Item A
Item B
Item C
[FINALIZAR_LISTA_MARCADORES]
```

Resultado:
```
• Item A
• Item B
• Item C
```

---

## 🔄 WORKFLOW DE USO

### **Passo 1: Ajustar Estrutura**

1. Abra `Sumario_Modelo.docx`
2. Adicione, remova ou reordene títulos conforme necessário
3. Mantenha numeração hierárquica consistente
4. Salve o arquivo

### **Passo 2: Ajustar Conteúdo**

1. Abra `Conteudo_Fonte.docx`
2. Edite os textos conforme necessário
3. Use marcadores `[INSERIR_TABELA_*]` onde tabelas devem aparecer
4. Use marcadores `[VARIAVEL]` onde valores dinâmicos devem aparecer
5. Salve o arquivo

### **Passo 3: Gerar Relatório**

```bash
python gerador_relatorio.py
```

O sistema irá:
- ✅ Ler a estrutura do sumário
- ✅ Ler o conteúdo mapeado
- ✅ Carregar dados do Excel
- ✅ Calcular variáveis dinâmicas
- ✅ Gerar documento final completo

---

## ⚠️ CUIDADOS IMPORTANTES

### **❌ NÃO FAÇA:**

- ❌ Não altere os marcadores `[MARCADOR]` (copie exatamente como está)
- ❌ Não adicione espaços dentro dos marcadores: `[ MARCADOR ]` ❌
- ❌ Não use marcadores inexistentes
- ❌ Não remova títulos do sumário se eles têm conteúdo

### **✅ FAÇA:**

- ✅ Mantenha títulos idênticos entre Sumario e Conteudo
- ✅ Use marcadores exatamente como documentado
- ✅ Teste após alterações estruturais
- ✅ Valide correspondência sumário ↔ conteúdo

---

## 🐛 TROUBLESHOOTING

### **Erro: "Título sem conteúdo"**

**Causa:** Título existe no sumário mas não no conteúdo.

**Solução:** Adicione seção correspondente no `Conteudo_Fonte.docx` ou remova do sumário.

### **Erro: "Conteúdo sem título no sumário"**

**Causa:** Seção existe no conteúdo mas não no sumário.

**Solução:** Adicione título correspondente no `Sumario_Modelo.docx` ou remova seção do conteúdo.

### **Variável não substituída (aparece [VARIAVEL] no relatório)**

**Causa:** Marcador de variável incorreto ou não implementado.

**Solução:** Verifique lista de marcadores válidos e corrija a digitação.

### **Tabela não aparece**

**Causa:** Marcador de tabela incorreto ou mal posicionado.

**Solução:** Verifique se o marcador está em uma linha separada, sem espaços extras.

---

## 📊 EXEMPLO COMPLETO

### **Sumario_Modelo.docx:**

```
1 INTRODUÇÃO
1.1 Contextualização

2 METAS APROVADAS
2.1 Total de Metas
```

### **Conteudo_Fonte.docx:**

```
1 INTRODUÇÃO

Este relatório apresenta os resultados de [NUMERO_METAS_2025] metas.

1.1 Contextualização

O processo de monitoramento tem como objetivo...

2 METAS APROVADAS

[INSERIR_TABELA_HISTORICA]

A tabela acima demonstra...

2.1 Total de Metas

#Destaque: [PERCENTUAL_VERDE]% das metas foram cumpridas!
```

### **Resultado Final:**

O sistema gera um documento Word com:
- ✅ Estrutura do sumário
- ✅ Textos formatados
- ✅ Tabelas dinâmicas inseridas
- ✅ Variáveis substituídas por valores reais
- ✅ Formatação especial aplicada

---

## 📞 SUPORTE

Para dúvidas ou problemas:

1. Consulte a documentação completa em `docs/PROPOSTA_IMPLEMENTACAO_TEMPLATES.md`
2. Verifique exemplos neste diretório
3. Valide templates antes de gerar relatório
4. Revise logs de processamento no console

---

**Versão:** 1.0  
**Última Atualização:** Dezembro/2025  
**Autor:** Sistema de Geração de Relatórios - TJMG
"""
    
    caminho = os.path.join(os.path.dirname(__file__), 'README_TEMPLATES.md')
    with open(caminho, 'w', encoding='utf-8') as f:
        f.write(readme_content)
    
    print(f"   ✅ Salvo em: {caminho}")
    
    return caminho


def main():
    """Função principal - cria todos os templates"""
    
    print("\n" + "="*70)
    print("📋 CRIADOR DE TEMPLATES - Sistema de Relatórios TJMG")
    print("="*70 + "\n")
    
    try:
        # Criar templates
        sumario = criar_sumario_modelo()
        conteudo = criar_conteudo_fonte()
        readme = criar_readme_templates()
        
        print("\n" + "="*70)
        print("✅ TEMPLATES CRIADOS COM SUCESSO!")
        print("="*70)
        print("\n📁 Arquivos gerados:")
        print(f"   1. {os.path.basename(sumario)}")
        print(f"   2. {os.path.basename(conteudo)}")
        print(f"   3. {os.path.basename(readme)}")
        print("\n💡 Próximos passos:")
        print("   • Abra os arquivos .docx e revise o conteúdo")
        print("   • Ajuste textos conforme necessário")
        print("   • Leia o README_TEMPLATES.md para instruções de uso")
        print("   • Execute o gerador de relatórios para testar")
        print("\n" + "="*70 + "\n")
        
    except Exception as e:
        print(f"\n❌ ERRO: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
