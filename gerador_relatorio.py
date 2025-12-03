"""
GERADOR DE RELATÓRIO DE METAS ESTRATÉGICAS - TJMG
Versão: 1.1 (Corrigido)
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from datetime import datetime
import os

# ============================================
# CONFIGURAÇÕES
# ============================================

class Config:
    """Configurações do relatório"""
    
    # Arquivo de entrada
    ARQUIVO_EXCEL = 'exports/teste_integração.xlsx'  # Nome do seu arquivo Excel
    NOME_ABA = None  # Nome da aba (ou deixe None para a primeira)
    
    # Arquivo de saída
    PASTA_SAIDA = 'relatorios_gerados'
    NOME_RELATORIO = 'Relatorio_Metas_Estrategicas'
    
    # Colunas do Excel (ajuste conforme seus dados)
    COLUNAS = {
        'MACRODESAFIO': 'Macrodesafio',
        'METAKEY': 'MetaKey',
        'ANO_META': 'Ano da Meta',
        'MES_APURACAO': 'mes_apuracao',
        'INDICADOR': 'Indicador',
        'UNIDADE_GESTORA': 'Unidade Gestora',
        'POLARIDADE': 'Polaridade',
        'VALOR_APURADO': 'Valor Apurado',
        'VALOR_META': 'Valor da Meta',
        'INICIATIVA': 'Iniciativa',
        'INFO_COMPLEMENTAR': 'Informação complementar texto'
    }
    
    # Cores (RGB) - Formato correto: (R, G, B)
    CORES = {
        'CABECALHO_MACRO': (239, 108, 33),      # Laranja forte
        'CABECALHO_TABELA': (250, 191, 143),    # Laranja claro
        'TEXTO_BRANCO': (255, 255, 255),
        'TEXTO_PRETO': (0, 0, 0),
        'BORDA_TABELA': (166, 166, 166),        # Cinza para bordas
        'INICIATIVA_FUNDO': (250, 191, 143)     # Laranja claro para iniciativa
    }
    
    # Formatação
    FONTE_PADRAO = 'Calibri'
    TAMANHO_TITULO = 18
    TAMANHO_SUBTITULO = 14
    TAMANHO_TEXTO = 11
    TAMANHO_TABELA = 9


# ============================================
# FUNÇÕES AUXILIARES
# ============================================

def criar_pasta_saida():
    """Cria pasta para salvar relatórios se não existir"""
    if not os.path.exists(Config.PASTA_SAIDA):
        os.makedirs(Config.PASTA_SAIDA)
        print(f"✅ Pasta '{Config.PASTA_SAIDA}' criada.")


def carregar_dados():
    """Carrega dados do Excel"""
    print(f"📂 Carregando dados de '{Config.ARQUIVO_EXCEL}'...")
    
    try:
        # Ler Excel
        if Config.NOME_ABA:
            df = pd.read_excel(Config.ARQUIVO_EXCEL, sheet_name=Config.NOME_ABA)
        else:
            df = pd.read_excel(Config.ARQUIVO_EXCEL)
        
        # Remover linhas completamente vazias
        df = df.dropna(how='all')
        
        print(f"✅ {len(df)} registros carregados.")
        return df
    
    except FileNotFoundError:
        print(f"❌ ERRO: Arquivo '{Config.ARQUIVO_EXCEL}' não encontrado!")
        print(f"   Certifique-se de que o arquivo está na mesma pasta que este script.")
        return None
    except Exception as e:
        print(f"❌ ERRO ao carregar dados: {e}")
        return None


def carregar_mapeamento_superintendencias():
    """Carrega mapeamento de metas para superintendências"""
    try:
        from meta_por_superintendencia import META_SUPERINTENDENCIA
        print(f"✅ Mapeamento de superintendências carregado ({len(META_SUPERINTENDENCIA)} metas).")
        return META_SUPERINTENDENCIA
    except ImportError:
        print("❌ ERRO: Módulo 'meta_por_superintendencia.py' não encontrado!")
        return {}
    except Exception as e:
        print(f"❌ ERRO ao carregar mapeamento: {e}")
        return {}


def extrair_codigo_meta(meta_texto):
    """Extrai código da meta (ex: 'TJMG 111 - Texto' -> 'TJMG 111')"""
    import re
    # Procurar padrão: letras maiúsculas seguidas de espaço e números
    match = re.match(r'^([A-Z]+\s+\d+)', str(meta_texto).strip())
    if match:
        return match.group(1)
    return None


def adicionar_coluna_superintendencia(df, mapeamento):
    """Adiciona coluna de superintendência ao DataFrame baseado no mapeamento"""
    print("\n🔍 Verificando mapeamento de superintendências...")
    
    def obter_superintendencia(meta):
        codigo = extrair_codigo_meta(meta)
        
        # Debug: imprimir primeiras 10 linhas para verificar
        if hasattr(obter_superintendencia, 'contador'):
            obter_superintendencia.contador += 1
        else:
            obter_superintendencia.contador = 1
        
        if obter_superintendencia.contador <= 10:
            resultado = mapeamento.get(codigo, 'SEM CLASSIFICAÇÃO') if codigo else 'SEM CLASSIFICAÇÃO'
            print(f"   Linha {obter_superintendencia.contador}: Meta='{str(meta)[:50]}...' | Código='{codigo}' | Superintendência='{resultado}'")
        
        if codigo:
            # Normalizar para maiúsculas para evitar divergências
            superintendencia = mapeamento.get(codigo, 'SEM CLASSIFICAÇÃO')
            return superintendencia.upper() if superintendencia != 'SEM CLASSIFICAÇÃO' else superintendencia
        return 'SEM CLASSIFICAÇÃO'
    
    df['Superintendência'] = df[Config.COLUNAS['METAKEY']].apply(obter_superintendencia)
    
    # Estatísticas de classificação
    classificacao_counts = df['Superintendência'].value_counts()
    print(f"\n📊 Estatísticas de classificação:")
    for super_nome, count in classificacao_counts.items():
        print(f"   {super_nome}: {count} registros")
    
    return df


def agrupar_por_superintendencia_e_macro(df):
    """Agrupa dados primeiro por Superintendência, depois por Macrodesafio"""
    print("📊 Agrupando dados por Superintendência e Macrodesafio...")
    
    # Importar ordem das superintendências do módulo
    from meta_por_superintendencia import ORDEM_SUPERINTENDENCIAS
    
    # Criar coluna auxiliar para ordenação de macrodesafio
    import re
    df['_ordem_macro'] = df[Config.COLUNAS['MACRODESAFIO']].apply(
        lambda x: int(re.match(r'^(\d+)', str(x)).group(1)) if pd.notna(x) and re.match(r'^(\d+)', str(x)) else 999
    )
    
    # Criar dicionário com ordem das superintendências
    ordem_dict = {super: idx for idx, super in enumerate(ORDEM_SUPERINTENDENCIAS)}
    df['_ordem_super'] = df['Superintendência'].map(lambda x: ordem_dict.get(x, 999))
    
    # Ordenar por superintendência e depois por macrodesafio
    df = df.sort_values(['_ordem_super', '_ordem_macro'])
    
    # Agrupar por superintendência
    grupos_super = {}
    for superintendencia in ORDEM_SUPERINTENDENCIAS:
        df_super = df[df['Superintendência'] == superintendencia]
        if len(df_super) > 0:
            # Dentro de cada superintendência, agrupar por macrodesafio
            grupos_macro = df_super.groupby(Config.COLUNAS['MACRODESAFIO'], sort=False)
            grupos_super[superintendencia] = list(grupos_macro)
    
    print(f"✅ {len(grupos_super)} Superintendências com dados encontradas.")
    return grupos_super


def agrupar_por_macrodesafio(df):
    """Agrupa dados por Macrodesafio"""
    print("📊 Agrupando dados por Macrodesafio...")
    
    col_macro = Config.COLUNAS['MACRODESAFIO']
    
    # Criar coluna auxiliar para ordenação numérica
    # Extrair número do início (ex: "1. Texto" -> 1)
    import re
    df['_ordem_macro'] = df[col_macro].apply(
        lambda x: int(re.match(r'^(\d+)', str(x)).group(1)) if pd.notna(x) and re.match(r'^(\d+)', str(x)) else 999
    )
    
    # Ordenar por número antes de agrupar
    df = df.sort_values('_ordem_macro')
    
    # Agrupar mantendo a ordem
    grupos = df.groupby(col_macro, sort=False)
    
    print(f"✅ {len(grupos)} Macrodesafios encontrados.")
    return grupos


def formatar_valor(valor):
    """Formata valor para exibição"""
    if pd.isna(valor) or valor == '':
        return '-'
    elif isinstance(valor, (int, float)):
        return f"{valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    else:
        return str(valor)


def rgb_to_hex(rgb_tuple):
    """Converte tupla RGB para hexadecimal"""
    return f"{rgb_tuple[0]:02x}{rgb_tuple[1]:02x}{rgb_tuple[2]:02x}"


def set_cell_background(cell, rgb_tuple):
    """Define cor de fundo de uma célula da tabela"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), rgb_to_hex(rgb_tuple))
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Define bordas para células da tabela"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remover bordas antigas se existirem
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if kwargs.get(edge):
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '2')  # 2 = 1/4pt (sz é em oitavos de ponto)
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), rgb_to_hex(Config.CORES['BORDA_TABELA']))
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)


def set_paragraph_background(paragraph, rgb_tuple):
    """Define cor de fundo de um parágrafo"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), rgb_to_hex(rgb_tuple))
    paragraph._element.get_or_add_pPr().append(shading_elm)


def set_keep_together(table):
    """Força que a tabela não quebre entre linhas"""
    tbl = table._element
    for row in tbl.findall('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        trPr = row.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row.insert(0, trPr)
        
        # Não permitir quebra dentro desta linha
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)


def set_keep_with_next(paragraph):
    """Força que o parágrafo/tabela permanece com o próximo elemento"""
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)


# ============================================
# FUNÇÕES DE CRIAÇÃO DO DOCUMENTO
# ============================================

def criar_documento(superintendencia='Presidência'):
    """Cria documento Word base com primeira página em retrato"""
    doc = Document()
    
    # === PRIMEIRA SEÇÃO: RETRATO (para tabela histórica) ===
    section_retrato = doc.sections[0]
    
    # Orientação retrato
    section_retrato.orientation = WD_ORIENT.PORTRAIT
    
    # Tamanho do papel A4 em retrato
    section_retrato.page_width = Cm(21.0)
    section_retrato.page_height = Cm(29.7)
    
    # Margens
    section_retrato.top_margin = Cm(2.5)
    section_retrato.bottom_margin = Cm(2.5)
    section_retrato.left_margin = Cm(2.0)
    section_retrato.right_margin = Cm(2.0)
    
    # Medianiz (gutter)
    section_retrato.gutter = Cm(0)
    
    # Cabeçalho e rodapé
    section_retrato.header_distance = Cm(1.5)
    section_retrato.footer_distance = Cm(1.5)
    
    # Cabeçalho da primeira página (retrato) - mais simples
    header_retrato = section_retrato.header
    
    # Limpar cabeçalho padrão
    for paragraph in list(header_retrato.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Título centralizado
    titulo_retrato = header_retrato.add_paragraph('Resultados do Monitoramento de Metas Estratégicas 2025')
    titulo_retrato.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_retrato.runs[0].font.size = Pt(14)
    titulo_retrato.runs[0].font.bold = True
    titulo_retrato.runs[0].font.name = Config.FONTE_PADRAO
    titulo_retrato.paragraph_format.space_after = Pt(6)
    
    # === RETORNAR DOCUMENTO (segunda seção paisagem será criada depois) ===
    return doc


def criar_secao_paisagem_inicial(doc, superintendencia='Presidência'):
    """Cria segunda seção em paisagem para as tabelas de metas"""
    
    # Adicionar quebra de seção para paisagem
    section_paisagem = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Configurar orientação paisagem
    section_paisagem.orientation = WD_ORIENT.LANDSCAPE
    
    # Tamanho do papel A4 em paisagem
    section_paisagem.page_width = Cm(29.7)
    section_paisagem.page_height = Cm(21.0)
    
    # Margens
    section_paisagem.top_margin = Cm(2.5)
    section_paisagem.bottom_margin = Cm(2.5)
    section_paisagem.left_margin = Cm(2.0)
    section_paisagem.right_margin = Cm(1.5)
    
    # Medianiz (gutter)
    section_paisagem.gutter = Cm(0)
    
    # Cabeçalho e rodapé
    section_paisagem.header_distance = Cm(1.05)
    section_paisagem.footer_distance = Cm(1.05)
    
    # Desvincular cabeçalho da seção anterior
    header = section_paisagem.header
    header.is_linked_to_previous = False
    
    # Limpar cabeçalho padrão
    for paragraph in list(header.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Título principal - negrito, 12pt, centralizado
    titulo = header.add_paragraph('Resultados do Monitoramento de Metas Estratégicas 2025')
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_format = titulo.runs[0]
    titulo_format.font.size = Pt(12)
    titulo_format.font.bold = True
    titulo_format.font.name = Config.FONTE_PADRAO
    titulo.paragraph_format.space_after = Pt(0)
    titulo.paragraph_format.space_before = Pt(0)
    
    # Subtítulo - normal, 11pt, centralizado
    subtitulo = header.add_paragraph('Relatório Técnico ao Comitê de Governança e Gestão Estratégica')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_format = subtitulo.runs[0]
    subtitulo_format.font.size = Pt(11)
    subtitulo_format.font.name = Config.FONTE_PADRAO
    subtitulo.paragraph_format.space_after = Pt(0)
    
    # Órgão - negrito, 12pt, laranja, centralizado (DINÂMICO)
    orgao = header.add_paragraph(superintendencia)
    orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    orgao_format = orgao.runs[0]
    orgao_format.font.size = Pt(12)
    orgao_format.font.bold = True
    orgao_format.font.name = Config.FONTE_PADRAO
    orgao_format.font.color.rgb = RGBColor(227, 108, 10)
    orgao.paragraph_format.space_after = Pt(6)
    
    # Adicionar conteúdo no rodapé
    footer = section_paisagem.footer
    footer.is_linked_to_previous = False
    
    # Limpar rodapé padrão
    for paragraph in list(footer.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Primeira linha do rodapé
    linha1 = footer.add_paragraph('Assessoria Técnica e Jurídica ao Planejamento e à Gestão Institucional - ASPLAG')
    linha1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    linha1_format = linha1.runs[0]
    linha1_format.font.size = Pt(9)
    linha1_format.font.name = Config.FONTE_PADRAO
    linha1.paragraph_format.space_after = Pt(0)
    
    # Segunda linha do rodapé
    linha2 = footer.add_paragraph('Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional - DEPLAG')
    linha2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    linha2_format = linha2.runs[0]
    linha2_format.font.size = Pt(9)
    linha2_format.font.name = Config.FONTE_PADRAO
    
    return section_paisagem


def adicionar_tabela_historica(doc):
    """
    Adiciona tabela histórica de metas aprovadas (4 anos mais recentes).
    Formato conforme imagem de referência com cabeçalho laranja.
    """
    from meta_por_superintendencia import atualizar_historico_com_ano_atual
    
    # Obter dados históricos (4 anos mais recentes)
    dados = atualizar_historico_com_ano_atual()
    historico = dados['historico_recente']
    anos = dados['anos_recentes']
    variacao = dados['variacao']
    primeiro_ano = anos[0]
    ultimo_ano = anos[-1]
    
    # Adicionar parágrafo de espaçamento
    doc.add_paragraph()
    
    # Criar tabela (4 linhas x 6 colunas)
    # Linha 1: Título (merged)
    # Linha 2: Cabeçalho (Ano, anos..., Variação)
    # Linha 3: Metas Nacionais
    # Linha 4: Metas Institucionais  
    # Linha 5: Total
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'
    
    # === LINHA 1: TÍTULO (MERGED) ===
    titulo_cells = table.rows[0].cells
    # Merge todas as células da primeira linha
    merged_cell = titulo_cells[0].merge(titulo_cells[5])
    
    # Adicionar texto do título
    titulo_paragraph = merged_cell.paragraphs[0]
    titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_paragraph.add_run(
        'TOTAL DE METAS APROVADAS PARA COMPOSIÇÃO DO PLANEJAMENTO ESTRATÉGICO\nINSTITUCIONAL DO TJMG'
    )
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor(255, 255, 255)
    titulo_run.font.name = Config.FONTE_PADRAO
    
    # Cor de fundo laranja
    shading = parse_xml(r'<w:shd {} w:fill="E36C0A"/>'.format(nsdecls('w')))
    merged_cell._element.get_or_add_tcPr().append(shading)
    
    # Alinhamento vertical centralizado
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Altura da linha do título: 0,7cm
    table.rows[0].height = Cm(0.7)
    
    # === LINHA 2: CABEÇALHO (Anos) ===
    header_row = table.rows[1]
    headers = ['Ano'] + [str(ano) for ano in anos] + [f'Variação\n{primeiro_ano} - {ultimo_ano}']
    
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        
        # Alinhamento: primeira coluna à esquerda, restante centralizado
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Formatação: tamanho 11, negrito
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
        
        # Cor de fundo RGB(250,191,143)
        shading = parse_xml(r'<w:shd {} w:fill="FABF8F"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
        
        # Alinhamento vertical centralizado
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Altura da linha do cabeçalho: 0,7cm
    table.rows[1].height = Cm(0.7)
    
    # === LINHA 3: METAS NACIONAIS ===
    row_nac = table.rows[2]
    row_nac.cells[0].text = 'Metas Nacionais'
    for idx, ano in enumerate(anos, start=1):
        row_nac.cells[idx].text = str(historico[ano]['nacionais'])
    row_nac.cells[5].text = f"{variacao['nacionais']:+}%"
    
    # Formatar células: tamanho 11, sem negrito
    for idx, cell in enumerate(row_nac.cells):
        paragraph = cell.paragraphs[0]
        
        # Alinhamento: primeira coluna à esquerda, restante centralizado
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.name = Config.FONTE_PADRAO
        
        # Alinhamento vertical centralizado
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Altura da linha: 1cm
    table.rows[2].height = Cm(1.0)
    
    # === LINHA 4: METAS INSTITUCIONAIS ===
    row_inst = table.rows[3]
    row_inst.cells[0].text = 'Metas Institucionais'
    for idx, ano in enumerate(anos, start=1):
        row_inst.cells[idx].text = str(historico[ano]['institucionais'])
    row_inst.cells[5].text = f"{variacao['institucionais']:+}%"
    
    # Formatar células: tamanho 11, sem negrito
    for idx, cell in enumerate(row_inst.cells):
        paragraph = cell.paragraphs[0]
        
        # Alinhamento: primeira coluna à esquerda, restante centralizado
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.name = Config.FONTE_PADRAO
        
        # Alinhamento vertical centralizado
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Altura da linha: 1cm
    table.rows[3].height = Cm(1.0)
    
    # === LINHA 5: TOTAL ===
    row_total = table.rows[4]
    row_total.cells[0].text = 'Total'
    for idx, ano in enumerate(anos, start=1):
        row_total.cells[idx].text = str(historico[ano]['total'])
    row_total.cells[5].text = f"{variacao['total']:+}%"
    
    # Formatar células: tamanho 11, negrito
    for idx, cell in enumerate(row_total.cells):
        paragraph = cell.paragraphs[0]
        
        # Alinhamento: primeira coluna à esquerda, restante centralizado
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
        
        # Cor de fundo RGB(64,64,64)
        shading = parse_xml(r'<w:shd {} w:fill="404040"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
        
        # Texto branco para linha Total
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Alinhamento vertical centralizado
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Altura da linha Total: 0,7cm
    table.rows[4].height = Cm(0.7)
    
    # Ajustar largura das colunas
    for row in table.rows:
        row.cells[0].width = Cm(4.5)  # Coluna "Ano"
        for i in range(1, 6):
            row.cells[i].width = Cm(2.5)  # Colunas dos anos e variação
    
    # Adicionar nota de rodapé
    nota = doc.add_paragraph()
    nota_run = nota.add_run(
        '*Destaca-se que a redução no número de metas nacionais no período decorreu de revisões propostas pelo '
        'Conselho Nacional de Justiça – CNJ, e aprovada pelos Tribunais Estaduais, em relação a seu compromisso com '
        'o Sistema de Justiça e a prestação jurisdicional.'
    )
    nota_run.font.size = Pt(8)
    nota_run.font.italic = True
    nota_run.font.name = Config.FONTE_PADRAO
    nota.paragraph_format.space_before = Pt(6)
    
    # Adicionar espaçamento após a tabela
    doc.add_paragraph()


def adicionar_tabela_macrodesafio(doc):
    """
    Adiciona tabela histórica de metas por macrodesafio (4 anos mais recentes).
    Formato conforme imagem de referência com cabeçalho laranja.
    """
    from meta_por_superintendencia import atualizar_historico_macrodesafio_com_ano_atual
    
    # Obter dados históricos (4 anos mais recentes)
    dados = atualizar_historico_macrodesafio_com_ano_atual()
    historico = dados['historico_recente']
    anos = dados['anos_recentes']
    macrodesafios = dados['macrodesafios']
    
    # Adicionar parágrafo de espaçamento
    doc.add_paragraph()
    
    # Criar tabela (linhas: título + cabeçalho + macrodesafios + total) x (colunas: macrodesafio + anos)
    num_linhas = 2 + len(macrodesafios) + 1  # título + cabeçalho + macros + total
    num_colunas = 1 + len(anos)  # macrodesafio + anos
    
    table = doc.add_table(rows=num_linhas, cols=num_colunas)
    table.style = 'Table Grid'
    
    # === LINHA 1: TÍTULO (MERGED) ===
    titulo_cells = table.rows[0].cells
    merged_cell = titulo_cells[0].merge(titulo_cells[-1])
    
    # Adicionar texto do título
    titulo_paragraph = merged_cell.paragraphs[0]
    titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_paragraph.add_run(
        f'TOTAL DE METAS POR MACRODESAFIO – COMPARATIVO {anos[0]} A {anos[-1]}'
    )
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor(255, 255, 255)
    titulo_run.font.name = Config.FONTE_PADRAO
    
    # Cor de fundo laranja
    shading = parse_xml(r'<w:shd {} w:fill="E36C0A"/>'.format(nsdecls('w')))
    merged_cell._element.get_or_add_tcPr().append(shading)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].height = Cm(0.8)
    
    # === LINHA 2: CABEÇALHO ===
    header_row = table.rows[1]
    headers = ['Macrodesafio'] + [str(ano) for ano in anos]
    
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        
        # Alinhamento: primeira coluna à esquerda, restante centralizado
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
        
        # Cor de fundo RGB(250,191,143)
        shading = parse_xml(r'<w:shd {} w:fill="FABF8F"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[1].height = Cm(0.8)
    
    # === LINHAS DE MACRODESAFIOS ===
    total_por_ano = {ano: 0 for ano in anos}
    
    for idx_macro, macrodesafio in enumerate(macrodesafios, start=2):
        row = table.rows[idx_macro]
        
        # Determinar cor de fundo alternada (baseado no índice do macrodesafio)
        idx_relativo = idx_macro - 2  # 0, 1, 2, ...
        if idx_relativo % 2 == 0:
            # Linhas pares: RGB(251,212,180)
            cor_fundo = "FBD4B4"
        else:
            # Linhas ímpares: RGB(255,255,255) - branco
            cor_fundo = "FFFFFF"
        
        # Primeira coluna: nome do macrodesafio
        row.cells[0].text = macrodesafio
        paragraph = row.cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.name = Config.FONTE_PADRAO
        
        # Aplicar cor de fundo
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[0]._element.get_or_add_tcPr().append(shading)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Demais colunas: valores por ano
        for idx_ano, ano in enumerate(anos, start=1):
            valor = historico[ano].get(macrodesafio, 0)
            total_por_ano[ano] += valor
            
            row.cells[idx_ano].text = str(valor)
            paragraph = row.cells[idx_ano].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.bold = False
                run.font.name = Config.FONTE_PADRAO
            
            # Aplicar mesma cor de fundo
            shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
            row.cells[idx_ano]._element.get_or_add_tcPr().append(shading)
            row.cells[idx_ano].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Altura da linha
        table.rows[idx_macro].height = Cm(0.7)
    
    # === LINHA TOTAL ===
    row_total_idx = num_linhas - 1
    row_total = table.rows[row_total_idx]
    
    row_total.cells[0].text = 'Total Geral'
    paragraph = row_total.cells[0].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
    
    # Cor de fundo e texto branco
    shading = parse_xml(r'<w:shd {} w:fill="404040"/>'.format(nsdecls('w')))
    row_total.cells[0]._element.get_or_add_tcPr().append(shading)
    
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    row_total.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Valores totais por ano
    for idx_ano, ano in enumerate(anos, start=1):
        row_total.cells[idx_ano].text = str(total_por_ano[ano])
        paragraph = row_total.cells[idx_ano].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
            run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Cor de fundo e texto branco
        shading = parse_xml(r'<w:shd {} w:fill="404040"/>'.format(nsdecls('w')))
        row_total.cells[idx_ano]._element.get_or_add_tcPr().append(shading)
        
        row_total.cells[idx_ano].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[row_total_idx].height = Cm(0.8)
    
    # Ajustar largura das colunas
    for row in table.rows:
        row.cells[0].width = Cm(12.0)  # Coluna "Macrodesafio" mais larga
        for i in range(1, num_colunas):
            row.cells[i].width = Cm(1.94)  # Colunas dos anos
    
    # Aplicar recuo negativo de -0.7cm à tabela
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr {}/> '.format(nsdecls('w')))
        tbl.insert(0, tblPr)
    
    tblInd = parse_xml(r'<w:tblInd {} w:w="-397" w:type="dxa"/>'.format(nsdecls('w')))
    tblPr.append(tblInd)
    
    # Adicionar espaçamento após a tabela
    doc.add_paragraph()


def adicionar_cabecalho_relatorio(doc):
    """Adiciona cabeçalho do relatório com tabelas históricas"""
    adicionar_tabela_historica(doc)
    adicionar_tabela_macrodesafio(doc)
    adicionar_tabela_resultado_monitoramento(doc)
    adicionar_tabela_metas_nacionais(doc)


def adicionar_tabela_metas_nacionais(doc):
    """
    Adiciona tabela de metas nacionais do CNJ.
    Mostra META NACIONAL, INSTÂNCIA e PERCENTUAL DE CUMPRIMENTO.
    Mescla células da mesma meta.
    """
    # Carregar dados do CNJ
    try:
        df_cnj = pd.read_excel('exports/resultados_cnj.xlsx')
    except FileNotFoundError:
        print("⚠️  Arquivo 'resultados_cnj.xlsx' não encontrado. Tabela de metas nacionais não será adicionada.")
        return
    
    # Adicionar parágrafo de espaçamento
    doc.add_paragraph()
    
    # Calcular número de linhas (cabeçalho + dados)
    num_linhas_dados = len(df_cnj)
    num_linhas = 1 + num_linhas_dados  # Apenas cabeçalho + dados
    
    # Criar tabela
    table = doc.add_table(rows=num_linhas, cols=3)
    table.style = 'Table Grid'
    
    # === LINHA 1: CABEÇALHO ===
    header_row = table.rows[0]
    headers = ['META NACIONAL', 'INSTÂNCIA', 'PERCENTUAL DE CUMPRIMENTO']
    
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = Config.FONTE_PADRAO
            run.font.color.rgb = RGBColor(255, 255, 255)  # Texto branco
        
        # Cor de fundo RGB(228,108,10)
        shading = parse_xml(r'<w:shd {} w:fill="E46C0A"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[0].height = Cm(0.8)
    
    # === PREPARAR DADOS PARA MESCLAGEM ===
    # Agrupar por Meta para identificar células a mesclar
    meta_ranges = {}  # {meta: (linha_inicial, linha_final, texto_meta)}
    meta_anterior = None
    linha_inicial_meta = 1  # Agora começa em 1 (linha após cabeçalho)
    
    for idx, (_, row_data) in enumerate(df_cnj.iterrows(), start=1):  # Começa em 1
        meta_completa = row_data.get('Meta', 'N/D')
        
        if meta_completa != meta_anterior:
            if meta_anterior is not None:
                # Salvar range da meta anterior (idx - 1 porque idx já é a nova meta)
                meta_ranges[meta_anterior] = (linha_inicial_meta, idx - 1)
            meta_anterior = meta_completa
            linha_inicial_meta = idx
    
    # Salvar a última meta
    if meta_anterior is not None:
        meta_ranges[meta_anterior] = (linha_inicial_meta, num_linhas_dados)
    
    # === LINHAS DE DADOS ===
    meta_atual = None
    for idx_linha, (_, row_data) in enumerate(df_cnj.iterrows(), start=1):  # Começa em 1
        row = table.rows[idx_linha]
        
        # Extrair dados
        meta_completa = row_data.get('Meta', 'N/D')
        categoria = row_data.get('Categoria', 'Total')
        resultado = row_data.get('Resultado', 'N/D')
        
        # Coluna 1: META NACIONAL (só preenche na primeira linha de cada meta)
        if meta_completa != meta_atual:
            meta_atual = meta_completa
            
            # Extrair apenas o número da meta
            import re
            match = re.search(r'Meta\s*(\d+)', meta_completa)
            numero_meta = match.group(1) if match else "?"
            
            # Montar texto: CNJ {numero} - {subtítulo} - {descrição}
            subtitulo = row_data.get('Subtítulo', '')
            descricao = row_data.get('Descrição', '')
            
            # Começar com "CNJ {numero}"
            texto_cnj = f"CNJ {numero_meta}"
            texto_complemento = ""
            
            # Adicionar subtítulo se não for N/D (removendo quebras de linha)
            if pd.notna(subtitulo) and subtitulo != 'N/D':
                subtitulo_limpo = str(subtitulo).replace('\n', ' ').replace('\r', ' ').strip()
                texto_complemento += f" - {subtitulo_limpo}"
            
            # Adicionar descrição se disponível (removendo quebras de linha)
            if pd.notna(descricao) and descricao != 'Descrição não encontrada':
                descricao_limpa = str(descricao).replace('\n', ' ').replace('\r', ' ').strip()
                texto_complemento += f" - {descricao_limpa}"
            
            # Limpar parágrafo e adicionar texto formatado
            row.cells[0].text = ""
            paragraph = row.cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Remover espaçamentos do parágrafo para deixar margens próximas ao texto
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Adicionar "CNJ {numero}" em negrito
            run_cnj = paragraph.add_run(texto_cnj)
            run_cnj.font.size = Pt(11)
            run_cnj.font.bold = True
            run_cnj.font.name = Config.FONTE_PADRAO
            
            # Adicionar resto do texto em fonte normal (sem quebra de linha)
            if texto_complemento:
                run_complemento = paragraph.add_run(texto_complemento)
                run_complemento.font.size = Pt(11)
                run_complemento.font.name = Config.FONTE_PADRAO
            
            # NÃO centralizar verticalmente, deixar no topo
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        else:
            # Limpa completamente o conteúdo das células que serão mescladas
            cell = row.cells[0]
            # Remover todos os parágrafos
            for paragraph in cell.paragraphs[1:]:
                p = paragraph._element
                p.getparent().remove(p)
            # Limpar o primeiro parágrafo
            cell.paragraphs[0].text = ""
        
        # Coluna 2: INSTÂNCIA
        # Remover duplicação (ex: "1º Grau1º Grau" -> "1º Grau")
        categoria_limpa = categoria
        if len(categoria) > 0 and len(categoria) % 2 == 0:
            metade = len(categoria) // 2
            primeira_metade = categoria[:metade]
            segunda_metade = categoria[metade:]
            if primeira_metade == segunda_metade:
                categoria_limpa = primeira_metade
        
        row.cells[1].text = categoria_limpa
        paragraph = row.cells[1].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = Config.FONTE_PADRAO
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Coluna 3: PERCENTUAL DE CUMPRIMENTO
        row.cells[2].text = str(resultado)
        paragraph = row.cells[2].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = Config.FONTE_PADRAO
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # === MESCLAR CÉLULAS DA COLUNA META NACIONAL E AJUSTAR ALTURAS ===
    for meta, (linha_ini, linha_fim) in meta_ranges.items():
        num_linhas_instancia = linha_fim - linha_ini + 1
        
        if linha_fim > linha_ini:  # Só mescla se houver mais de uma linha
            cell_inicial = table.rows[linha_ini].cells[0]
            cell_final = table.rows[linha_fim].cells[0]
            cell_inicial.merge(cell_final)
            
            # Remover parágrafos vazios extras da célula mesclada
            paragraphs_to_remove = []
            for i, paragraph in enumerate(cell_inicial.paragraphs):
                if i > 0:  # Manter apenas o primeiro parágrafo
                    paragraphs_to_remove.append(paragraph)
            
            for paragraph in paragraphs_to_remove:
                p = paragraph._element
                p.getparent().remove(p)
            
            # Centralizar verticalmente o texto na célula mesclada
            cell_inicial.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # === AJUSTAR ALTURA DAS LINHAS ===
        # Calcular altura mínima necessária baseada nas linhas de Instância
        altura_minima_instancia = Cm(num_linhas_instancia * 0.7)
        
        # Definir altura para cada linha do range
        for idx_linha in range(linha_ini, linha_fim + 1):
            # A altura de cada linha individual deve ser proporcional
            table.rows[idx_linha].height = Cm(0.7)
    
    # === DEFINIR LARGURAS DAS COLUNAS ===
    # Forçar larguras específicas das colunas
    table.columns[0].width = Cm(10.5)  # META NACIONAL
    table.columns[1].width = Cm(3.75)  # INSTÂNCIA
    table.columns[2].width = Cm(3.75)  # PERCENTUAL DE CUMPRIMENTO
    
    # Garantir larguras também em cada célula
    for row in table.rows:
        row.cells[0].width = Cm(10.5)  # META NACIONAL
        row.cells[1].width = Cm(3.75)  # INSTÂNCIA
        row.cells[2].width = Cm(3.75)  # PERCENTUAL DE CUMPRIMENTO
    
    # === APLICAR BORDAS PERSONALIZADAS ===
    # Remover todas as bordas e adicionar apenas borda inferior de 1pt
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            
            # Remover bordas existentes se houver
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)
            
            # Criar novas bordas
            tcBorders = OxmlElement('w:tcBorders')
            
            # Borda superior: none
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'none')
            top.set(qn('w:sz'), '0')
            tcBorders.append(top)
            
            # Borda esquerda: none
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'none')
            left.set(qn('w:sz'), '0')
            tcBorders.append(left)
            
            # Borda inferior: single 1pt
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')  # 1pt = 8 eighths of a point
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)
            
            # Borda direita: none
            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'none')
            right.set(qn('w:sz'), '0')
            tcBorders.append(right)
            
            tcPr.append(tcBorders)
    
    # === SEM RECUO PARA ESTA TABELA ===
    # Tabela de metas nacionais não tem recuo aplicado
    
    print(f"✅ Tabela de Metas Nacionais adicionada ({num_linhas_dados} linhas)")


def adicionar_tabela_resultado_monitoramento(doc):
    """
    Adiciona tabela de resultado do monitoramento de metas.
    Mostra distribuição por faixa de cumprimento.
    Une dados de metas institucionais (TJMG) e metas nacionais (CNJ).
    """
    from meta_por_superintendencia import calcular_cumprimento_metas_cnj
    
    print("\n📊 Calculando resultado do monitoramento...")
    
    # === CALCULAR CUMPRIMENTO METAS CNJ ===
    metas_cnj_cumprimento = calcular_cumprimento_metas_cnj()
    
    # === CALCULAR CUMPRIMENTO METAS INSTITUCIONAIS ===
    try:
        df_tjmg = pd.read_excel('exports/teste_integração.xlsx')
        
        # Converter percentual para float
        def converter_percentual(val):
            if pd.isna(val):
                return 0.0
            val_str = str(val).replace('%', '').replace(',', '.').strip()
            try:
                return float(val_str)
            except:
                return 0.0
        
        df_tjmg['Resultado_Float'] = df_tjmg['Valor Apurado'].apply(converter_percentual)
        
        # Para metas institucionais, usar valor apurado diretamente
        metas_tjmg_cumprimento = {}
        for _, row in df_tjmg.iterrows():
            meta_key = row['MetaKey']
            # Extrair código (TJMG XXX)
            import re
            match = re.search(r'^(TJMG\s+\d+)', meta_key)
            if match:
                codigo = match.group(1)
                cumprimento = row['Resultado_Float']
                metas_tjmg_cumprimento[codigo] = {'cumprimento': cumprimento, 'categorias': 1}
    
    except Exception as e:
        print(f"⚠️  Erro ao processar metas TJMG: {e}")
        metas_tjmg_cumprimento = {}
    
    # === UNIR TODOS OS CUMPRIMENTOS ===
    todos_cumprimentos = {}
    todos_cumprimentos.update(metas_cnj_cumprimento)
    todos_cumprimentos.update(metas_tjmg_cumprimento)
    
    # === CLASSIFICAR POR FAIXA ===
    maior_100 = 0
    entre_70_100 = 0
    abaixo_70 = 0
    sem_apuracao = 0
    
    for meta, dados in todos_cumprimentos.items():
        cumprimento = dados['cumprimento']
        
        if cumprimento >= 100:
            maior_100 += 1
        elif cumprimento >= 70:
            entre_70_100 += 1
        elif cumprimento > 0:
            abaixo_70 += 1
        else:
            sem_apuracao += 1
    
    total = len(todos_cumprimentos)
    
    print(f"✅ Processadas {total} metas ({len(metas_cnj_cumprimento)} CNJ + {len(metas_tjmg_cumprimento)} TJMG)")
    
    # Adicionar parágrafo de espaçamento
    doc.add_paragraph()
    
    # Criar tabela (linhas: título + 4 faixas + total) x 3 colunas (SEM cabeçalho)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    # === LINHA 1: TÍTULO (MERGED) ===
    titulo_cells = table.rows[0].cells
    merged_cell = titulo_cells[0].merge(titulo_cells[-1])
    
    titulo_paragraph = merged_cell.paragraphs[0]
    titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_paragraph.add_run('RESULTADO DO MONITORAMENTO DAS METAS ESTRATÉGICAS')
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor(255, 255, 255)
    titulo_run.font.name = Config.FONTE_PADRAO
    
    # Cor de fundo laranja
    shading = parse_xml(r'<w:shd {} w:fill="E36C0A"/>'.format(nsdecls('w')))
    merged_cell._element.get_or_add_tcPr().append(shading)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].height = Cm(0.8)
    
    # === LINHAS DE DADOS (começam na linha 1, índice ajustado) ===
    faixas = [
        ('Metas com resultado maior ou igual 100%', maior_100),
        ('Metas com resultado entre 70% e 100%', entre_70_100),
        ('Metas com resultado abaixo de 70%', abaixo_70),
        ('Metas sem apuração até outubro/2025', sem_apuracao)
    ]
    
    for idx_faixa, (descricao, quantidade) in enumerate(faixas, start=1):
        row = table.rows[idx_faixa]
        
        # Definir cor de fundo alternada
        idx_relativo = idx_faixa - 1  # 0, 1, 2, 3
        if idx_relativo % 2 == 0:
            cor_fundo = "D9D9D9"  # RGB(217,217,217) - Primeira linha
        else:
            cor_fundo = "FFFFFF"  # RGB(255,255,255) - Segunda linha (branco)
        
        # Coluna 1: Descrição
        row.cells[0].text = descricao
        paragraph = row.cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[0]._element.get_or_add_tcPr().append(shading)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Coluna 2: Quantidade (NEGRITO)
        row.cells[1].text = str(quantidade)
        paragraph = row.cells[1].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True  # NEGRITO
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[1]._element.get_or_add_tcPr().append(shading)
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Coluna 3: Percentual (NEGRITO)
        percentual = round((quantidade / total * 100)) if total > 0 else 0
        row.cells[2].text = f"{percentual}%"
        paragraph = row.cells[2].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True  # NEGRITO
            run.font.name = Config.FONTE_PADRAO
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_fundo))
        row.cells[2]._element.get_or_add_tcPr().append(shading)
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Altura da linha
        table.rows[idx_faixa].height = Cm(1.0)
    
    # === LINHA TOTAL ===
    row_total = table.rows[5]
    
    row_total.cells[0].text = 'Total'
    paragraph = row_total.cells[0].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    shading = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
    row_total.cells[0]._element.get_or_add_tcPr().append(shading)
    row_total.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Coluna quantidade (NEGRITO)
    row_total.cells[1].text = str(total)
    paragraph = row_total.cells[1].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    shading = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
    row_total.cells[1]._element.get_or_add_tcPr().append(shading)
    row_total.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Coluna percentual (NEGRITO)
    row_total.cells[2].text = '100%'
    paragraph = row_total.cells[2].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = Config.FONTE_PADRAO
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    shading = parse_xml(r'<w:shd {} w:fill="595959"/>'.format(nsdecls('w')))
    row_total.cells[2]._element.get_or_add_tcPr().append(shading)
    row_total.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    table.rows[5].height = Cm(1.0)
    
    # Ajustar largura das colunas
    for row in table.rows:
        row.cells[0].width = Cm(8.1)    # Coluna descrição
        row.cells[1].width = Cm(2.75)   # Coluna quantidade
        row.cells[2].width = Cm(4.0)    # Coluna percentual
    
    # === APLICAR RECUO DE +1,5cm NA TABELA ===
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Adicionar recuo de +1,5cm (positivo = para direita)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(1.5 * 567)))  # Converter cm para twips (1cm = 567 twips)
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Adicionar espaçamento após a tabela
    doc.add_paragraph()


def adicionar_nova_secao_superintendencia(doc, superintendencia, primeira=False):
    """Adiciona nova seção com cabeçalho personalizado para superintendência"""
    
    if not primeira:
        # Adicionar quebra de seção (cria nova seção automaticamente)
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        # Usar a última seção existente
        new_section = doc.sections[-1]
    
    # Configurar a nova seção
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)
    new_section.left_margin = Cm(2.0)
    new_section.right_margin = Cm(1.5)
    new_section.gutter = Cm(0)
    new_section.header_distance = Cm(1.05)
    new_section.footer_distance = Cm(1.05)
    
    # IMPORTANTE: Desvincular cabeçalho e rodapé da seção anterior
    new_section.different_first_page_header_footer = False
    
    # Configurar cabeçalho da nova seção
    header = new_section.header
    
    # Desvincular do cabeçalho anterior
    header.is_linked_to_previous = False
    
    # Limpar cabeçalho padrão
    for paragraph in list(header.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Título principal
    titulo = header.add_paragraph('Resultados do Monitoramento de Metas Estratégicas 2025')
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_format = titulo.runs[0]
    titulo_format.font.size = Pt(12)
    titulo_format.font.bold = True
    titulo_format.font.name = Config.FONTE_PADRAO
    titulo.paragraph_format.space_after = Pt(0)
    titulo.paragraph_format.space_before = Pt(0)
    
    # Subtítulo
    subtitulo = header.add_paragraph('Relatório Técnico ao Comitê de Governança e Gestão Estratégica')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_format = subtitulo.runs[0]
    subtitulo_format.font.size = Pt(11)
    subtitulo_format.font.name = Config.FONTE_PADRAO
    subtitulo.paragraph_format.space_after = Pt(0)
    
    # Superintendência (DINÂMICO)
    orgao = header.add_paragraph(superintendencia)
    orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    orgao_format = orgao.runs[0]
    orgao_format.font.size = Pt(12)
    orgao_format.font.bold = True
    orgao_format.font.name = Config.FONTE_PADRAO
    orgao_format.font.color.rgb = RGBColor(227, 108, 10)
    orgao.paragraph_format.space_after = Pt(6)
    
    # Configurar rodapé da nova seção
    footer = new_section.footer
    
    # Desvincular do rodapé anterior
    footer.is_linked_to_previous = False
    
    # Limpar rodapé padrão
    for paragraph in list(footer.paragraphs):
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    # Primeira linha do rodapé
    linha1 = footer.add_paragraph('Assessoria Técnica e Jurídica ao Planejamento e à Gestão Institucional - ASPLAG')
    linha1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    linha1_format = linha1.runs[0]
    linha1_format.font.size = Pt(9)
    linha1_format.font.name = Config.FONTE_PADRAO
    linha1.paragraph_format.space_after = Pt(0)
    
    # Segunda linha do rodapé
    linha2 = footer.add_paragraph('Diretoria Executiva de Planejamento Orçamentário e Qualidade na Gestão Institucional - DEPLAG')
    linha2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    linha2_format = linha2.runs[0]
    linha2_format.font.size = Pt(9)
    linha2_format.font.name = Config.FONTE_PADRAO
    linha2.paragraph_format.space_after = Pt(0)


def adicionar_secao_macrodesafio(doc, macrodesafio, df_grupo, primeira_secao=False):
    """Adiciona seção de um Macrodesafio"""
    
    # Quebra de página entre macrodesafios (exceto no primeiro)
    if not primeira_secao:
        doc.add_page_break()
    
    def adicionar_cabecalho_macro():
        """Função auxiliar para adicionar cabeçalho do macrodesafio"""
        # Título do Macrodesafio usando tabela para controlar largura
        titulo_tabela = doc.add_table(rows=1, cols=1)
        titulo_tabela.style = None
        titulo_cell = titulo_tabela.rows[0].cells[0]
        titulo_cell.text = macrodesafio.upper()
        
        # Formatação do título
        for paragraph in titulo_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(10)  # Espaço antes
            paragraph.paragraph_format.space_after = Pt(10)   # Espaço depois
            paragraph.paragraph_format.line_spacing = 1.0     # Espaçamento entre linhas
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = Config.FONTE_PADRAO
                run.font.color.rgb = RGBColor(*Config.CORES['TEXTO_BRANCO'])
                run.font.bold = True
        
        # Cor de fundo e alinhamento vertical
        set_cell_background(titulo_cell, Config.CORES['CABECALHO_MACRO'])
        from docx.enum.table import WD_ALIGN_VERTICAL
        titulo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Definir altura da célula: 1,03cm
        titulo_tabela.rows[0].height = Cm(1.03)
        
        # Remover bordas
        set_cell_border(titulo_cell, top=False, bottom=False, left=False, right=False)
        
        # Definir largura igual à soma das colunas da tabela de dados
        # Largura total: 28,5cm (de -1cm até 27,5cm na régua)
        largura_total_tabela = Cm(28.5)
        titulo_tabela.rows[0].cells[0].width = largura_total_tabela
        
        # Adicionar recuo negativo de 1cm
        tbl = titulo_tabela._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '-567')  # -1cm em twips (1cm = 567 twips)
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        # Marcar que já foi executado uma vez
        adicionar_cabecalho_macro._ja_executado = True
    
    # Adicionar cabeçalho inicial
    adicionar_cabecalho_macro()
    
    # Processar cada indicador individualmente com controle de quebra
    precisa_cabecalho = False
    incluir_subcabecalho = True  # Primeiro sempre tem subcabeçalho
    
    for idx, (_, row) in enumerate(df_grupo.iterrows()):
        # Se precisa cabeçalho (após quebra de página), adicionar
        if precisa_cabecalho:
            doc.add_page_break()
            adicionar_cabecalho_macro()
            incluir_subcabecalho = True  # Após quebra, sempre incluir subcabeçalho
            precisa_cabecalho = False
        
        # Criar tabela para este indicador
        largura_total = adicionar_tabela_indicador(doc, row, incluir_cabecalho=incluir_subcabecalho)
        incluir_subcabecalho = False  # Próximos não terão subcabeçalho (a menos que haja quebra)
        
        # Adicionar informação complementar se existir
        info_complementar = row.get(Config.COLUNAS['INFO_COMPLEMENTAR'], '')
        tem_situacao = pd.notna(info_complementar) and str(info_complementar).strip() != ''
        
        if tem_situacao:
            texto_situacao = str(info_complementar)
            
            # Criar tabela de célula única para situação (usado tanto para texto curto quanto blocos longos)
            
            # Criar tabela de célula única para situação
            situacao_tabela = doc.add_table(rows=1, cols=1)
            situacao_tabela.style = None
            situacao_cell = situacao_tabela.rows[0].cells[0]
            
            # Processar texto e adicionar formatação
            # Dividir o texto em partes: antes dos marcadores e itens com marcadores
            import re
            
            # Separar por "•" mantendo o marcador
            partes = re.split(r'(•)', texto_situacao)
            
            primeiro_paragrafo = True
            paragrafo_atual = None
            texto_acumulado = ""
            
            for i, parte in enumerate(partes):
                if parte == '•':
                    # Finalizar parágrafo anterior se houver texto acumulado
                    if texto_acumulado.strip():
                        if paragrafo_atual is None:
                            paragrafo_atual = situacao_cell.paragraphs[0] if primeiro_paragrafo and len(situacao_cell.paragraphs) > 0 else situacao_cell.add_paragraph()
                            primeiro_paragrafo = False
                        
                        if paragrafo_atual.text == '':
                            run_bold = paragrafo_atual.add_run('Situação: ')
                            run_bold.font.bold = True
                            run_bold.font.size = Pt(10)
                            run_bold.font.name = Config.FONTE_PADRAO
                        
                        run = paragrafo_atual.add_run(texto_acumulado)
                        run.font.size = Pt(10)
                        run.font.name = Config.FONTE_PADRAO
                        
                        paragrafo_atual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragrafo_atual.paragraph_format.space_before = Pt(0)
                        paragrafo_atual.paragraph_format.space_after = Pt(0)
                        
                        texto_acumulado = ""
                        paragrafo_atual = None
                    
                    # Iniciar novo parágrafo com marcador
                    paragrafo_atual = situacao_cell.add_paragraph()
                    primeiro_paragrafo = False
                    
                    run_marcador = paragrafo_atual.add_run('•')
                    run_marcador.font.size = Pt(10)
                    run_marcador.font.name = Config.FONTE_PADRAO
                    
                    paragrafo_atual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragrafo_atual.paragraph_format.space_before = Pt(0)
                    paragrafo_atual.paragraph_format.space_after = Pt(0)
                else:
                    texto_acumulado += parte
            
            # Adicionar último texto acumulado
            if texto_acumulado.strip():
                if paragrafo_atual is None:
                    if len(situacao_cell.paragraphs) > 0 and primeiro_paragrafo:
                        paragrafo_atual = situacao_cell.paragraphs[0]
                    else:
                        paragrafo_atual = situacao_cell.add_paragraph()
                
                # Adicionar "Situação:" apenas se for o primeiro parágrafo e estiver vazio
                if len(situacao_cell.paragraphs) == 1 and paragrafo_atual.text == '':
                    run_bold = paragrafo_atual.add_run('Situação: ')
                    run_bold.font.bold = True
                    run_bold.font.size = Pt(10)
                    run_bold.font.name = Config.FONTE_PADRAO
                
                run = paragrafo_atual.add_run(texto_acumulado)
                run.font.size = Pt(10)
                run.font.name = Config.FONTE_PADRAO
                
                paragrafo_atual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragrafo_atual.paragraph_format.space_before = Pt(0)
                paragrafo_atual.paragraph_format.space_after = Pt(0)
            
            # Formatação da célula
            from docx.enum.table import WD_ALIGN_VERTICAL
            situacao_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Definir mesma largura da tabela principal
            situacao_tabela.rows[0].cells[0].width = largura_total
            
            # Adicionar bordas
            set_cell_border(situacao_cell, top=True, bottom=True, left=True, right=True)
            
            # Adicionar recuo negativo de 1cm à tabela de situação
            tbl_situacao = situacao_tabela._element
            tblPr_situacao = tbl_situacao.tblPr
            if tblPr_situacao is None:
                tblPr_situacao = OxmlElement('w:tblPr')
                tbl_situacao.insert(0, tblPr_situacao)
            
            tblInd_situacao = OxmlElement('w:tblInd')
            tblInd_situacao.set(qn('w:w'), '-567')  # -1cm em twips
            tblInd_situacao.set(qn('w:type'), 'dxa')
            tblPr_situacao.append(tblInd_situacao)
            
            # Permitir que a tabela de situação quebre entre páginas se necessário
            # Configurar para permitir quebra de linha nas células
            for row in situacao_tabela.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    # Remover restrição cantSplit se existir
                    cantSplit_elements = tcPr.findall(qn('w:cantSplit'))
                    for cs in cantSplit_elements:
                        tcPr.remove(cs)
        
        # Verificar se próximo indicador precisa de nova página
        # (lógica simplès: se tem situação longa, próximo item começa com cabeçalho)
        if tem_situacao and len(str(info_complementar)) > 800:
            precisa_cabecalho = True


def adicionar_tabela_indicador(doc, row, incluir_cabecalho=True):
    """Adiciona tabela com um único indicador e retorna a largura total"""
    
    # Criar tabela (cabeçalho opcional + 1 linha de dados)
    num_linhas = 2 if incluir_cabecalho else 1
    num_colunas = 7  # Incluindo coluna de indicador visual
    tabela = doc.add_table(rows=num_linhas, cols=num_colunas)
    tabela.style = None  # Remover estilo padrão para controle total
    
    # Cabeçalhos (se necessário)
    if incluir_cabecalho:
        headers = ['INDICADOR', 'META', 'UNIDADE RESPONSÁVEL', 'POLARIDADE', 'RESULTADO APURADO', 'INICIATIVA(S)']
        header_cells = tabela.rows[0].cells
        
        # Mesclar células 4 e 5 (RESULTADO APURADO e ●)
        # No python-docx, mesclamos referenciando células
        header_cells[4].merge(header_cells[5])
        
        for i, header in enumerate(headers):
            # Pular índice 5 pois foi mesclado com 4
            if i >= 5:
                cell_index = i + 1
            else:
                cell_index = i
            
            cell = header_cells[cell_index]
            cell.text = header
            
            # Formatação do cabeçalho
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centro horizontal
                paragraph.paragraph_format.space_before = Pt(0)  # Espaço antes
                paragraph.paragraph_format.space_after = Pt(0)   # Espaço depois
                paragraph.paragraph_format.line_spacing = 1.0     # Espaçamento entre linhas
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = Config.FONTE_PADRAO
                    run.font.color.rgb = RGBColor(*Config.CORES['TEXTO_PRETO'])
            
            from docx.enum.table import WD_ALIGN_VERTICAL
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Centro vertical
            
            # Cor de fundo
            set_cell_background(cell, Config.CORES['CABECALHO_TABELA'])
            
            # Adicionar bordas brancas
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
        
        # Definir altura da linha do cabeçalho: 1,03cm
        tabela.rows[0].height = Cm(1.03)
    
    # Preencher dados (última linha da tabela)
    linha_dados = 1 if incluir_cabecalho else 0
    cells = tabela.rows[linha_dados].cells
    
    # Dados
    dados = [
        formatar_valor(row.get(Config.COLUNAS['INDICADOR'], '-')),
        formatar_valor(row.get(Config.COLUNAS['METAKEY'], '-')),
        formatar_valor(row.get(Config.COLUNAS['UNIDADE_GESTORA'], '-')),
        formatar_valor(row.get(Config.COLUNAS['POLARIDADE'], '-')),
        formatar_valor(row.get(Config.COLUNAS['VALOR_APURADO'], '-')),
        '●',  # Indicador visual
        formatar_valor(row.get(Config.COLUNAS['INICIATIVA'], '-'))
    ]
    
    # Determinar cor do indicador visual baseado no atingimento da meta
    valor_apurado = row.get(Config.COLUNAS['VALOR_APURADO'], None)
    valor_meta = row.get(Config.COLUNAS['VALOR_META'], None)
    
    # Debug: imprimir valores para primeira meta
    meta_key = row.get(Config.COLUNAS['METAKEY'], '')
    if 'TJMG 111' in str(meta_key):
        print(f"\n🔍 DEBUG TJMG 111:")
        print(f"   Meta Key: {meta_key}")
        print(f"   Valor Apurado (raw): {repr(valor_apurado)} | Tipo: {type(valor_apurado)}")
        print(f"   Valor Meta (raw): {repr(valor_meta)} | Tipo: {type(valor_meta)}")
        print(f"   pd.notna(valor_apurado): {pd.notna(valor_apurado)}")
        print(f"   pd.notna(valor_meta): {pd.notna(valor_meta)}")
    
    # Verde se atingiu a meta, vermelho caso contrário
    cor_indicador = RGBColor(0, 255, 0)  # Verde (padrão)
    if pd.notna(valor_apurado) and pd.notna(valor_meta):
        try:
            # Converter valor apurado (pode estar como string com vírgula)
            if isinstance(valor_apurado, str):
                val_apurado_float = float(valor_apurado.replace('.', '').replace(',', '.'))
            else:
                val_apurado_float = float(valor_apurado)
            
            # Converter valor da meta
            val_meta_float = float(valor_meta)
            
            if 'TJMG 111' in str(meta_key):
                print(f"   Valor Apurado (float): {val_apurado_float}")
                print(f"   Valor Meta (float): {val_meta_float}")
                print(f"   Comparação: {val_apurado_float} < {val_meta_float} = {val_apurado_float < val_meta_float}")
            
            if val_apurado_float < val_meta_float:
                cor_indicador = RGBColor(255, 0, 0)  # Vermelho
                if 'TJMG 111' in str(meta_key):
                    print(f"   ✅ Cor definida: VERMELHO")
            else:
                if 'TJMG 111' in str(meta_key):
                    print(f"   ✅ Cor definida: VERDE")
        except (ValueError, TypeError) as e:
            if 'TJMG 111' in str(meta_key):
                print(f"   ❌ Erro na conversão: {e}")
            pass  # Mantém verde se não conseguir converter
    
    for i, dado in enumerate(dados):
        cell = cells[i]
        
        # Tratamento especial para coluna Indicador (i == 0)
        if i == 0:
            # Limpar célula
            cell.text = ''
            paragraph = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
            
            # Separar número do texto (ex: "1.14 - Texto" ou "1.14 Texto")
            import re
            match = re.match(r'^(\d+\.?\d*)\s*[-–]?\s*(.*)$', dado)
            
            if match:
                numero = match.group(1)
                texto_resto = match.group(2)
                
                # Adicionar número em negrito
                run_negrito = paragraph.add_run(numero)
                run_negrito.font.bold = True
                run_negrito.font.size = Pt(Config.TAMANHO_TABELA)
                run_negrito.font.name = Config.FONTE_PADRAO
                
                # Adicionar o resto do texto normalmente
                if texto_resto:
                    run_normal = paragraph.add_run(' - ' + texto_resto)
                    run_normal.font.size = Pt(Config.TAMANHO_TABELA)
                    run_normal.font.name = Config.FONTE_PADRAO
            else:
                # Se não encontrar padrão, adicionar texto normal
                run = paragraph.add_run(dado)
                run.font.size = Pt(Config.TAMANHO_TABELA)
                run.font.name = Config.FONTE_PADRAO
            
            # Alinhamento
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        
        # Tratamento especial para coluna Meta (i == 1)
        elif i == 1:
            # Limpar célula
            cell.text = ''
            paragraph = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
            
            # Separar código da meta do texto (ex: "TJMG 111 - Texto" ou "TJMG 111")
            import re
            match = re.match(r'^([A-Z]+\s+\d+)\s*[-–]?\s*(.*)$', dado)
            
            if match:
                codigo = match.group(1)
                texto_resto = match.group(2)
                
                # Adicionar código em negrito
                run_negrito = paragraph.add_run(codigo)
                run_negrito.font.bold = True
                run_negrito.font.size = Pt(Config.TAMANHO_TABELA)
                run_negrito.font.name = Config.FONTE_PADRAO
                
                # Adicionar o resto do texto normalmente
                if texto_resto:
                    run_normal = paragraph.add_run(' - ' + texto_resto)
                    run_normal.font.size = Pt(Config.TAMANHO_TABELA)
                    run_normal.font.name = Config.FONTE_PADRAO
            else:
                # Se não encontrar padrão, adicionar texto normal
                run = paragraph.add_run(dado)
                run.font.size = Pt(Config.TAMANHO_TABELA)
                run.font.name = Config.FONTE_PADRAO
            
            # Alinhamento
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        
        else:
            # Outras colunas: formatação normal
            cell.text = dado
        
        # Formatação do texto (para colunas que não são Indicador nem Meta)
        if i not in [0, 1]:
            for paragraph in cell.paragraphs:
                # Alinhamento - centralizar todos os dados
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Remover espaçamentos
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            
            for run in paragraph.runs:
                run.font.size = Pt(Config.TAMANHO_TABELA)
                run.font.name = Config.FONTE_PADRAO
                
                # Destacar resultado apurado
                if i == 4:  # Coluna de Resultado Apurado
                    run.font.bold = True
                
                # Cor do indicador visual (verde ou vermelho baseado no atingimento da meta)
                if i == 5:  # Coluna ●
                    run.font.color.rgb = cor_indicador
                    run.font.size = Pt(16)
        
        # Remover cor de fundo para coluna Iniciativa (agora é a coluna 6)
        # Não aplicar nenhuma cor de fundo
        
        from docx.enum.table import WD_ALIGN_VERTICAL
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Centro vertical
        
        # Adicionar bordas brancas
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)
    
    # Ajustar largura das colunas (em centímetros)
    # Largura mínima de 1,3cm para todas
    larguras = [
        None,        # INDICADOR - não alterável (automática)
        Cm(9.09),    # META
        Cm(2.28),    # UNIDADE RESPONSÁVEL
        Cm(2.1),     # POLARIDADE
        Cm(1.97),    # RESULTADO APURADO
        Cm(0.72),    # ●
        Cm(2.18)     # INICIATIVA(S)
    ]
    
    # Calcular largura do Indicador baseado no espaço restante
    # Largura total desejada: 28,5cm (de -1cm até 27,5cm na régua)
    largura_total = Cm(28.5)
    largura_outras = sum([l for l in larguras if l is not None])
    largura_indicador = largura_total - largura_outras
    larguras[0] = max(largura_indicador, Cm(1.3))  # Mínimo de 1,3cm
    
    for row_table in tabela.rows:
        for idx, width in enumerate(larguras):
            if width:
                row_table.cells[idx].width = width
    
    # Calcular largura total da tabela para retornar
    largura_total_tabela = sum([l for l in larguras if l is not None])
    
    # Configurar propriedades da tabela para controle de quebra de página
    # Manter linhas da tabela juntas quando possível
    tbl = tabela._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Adicionar recuo negativo de 1cm
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '-567')  # -1cm em twips
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Permitir que a tabela quebre entre linhas se necessário
    tblPrEx = OxmlElement('w:tblPrEx')
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), '0')  # Permitir quebra
    tblPrEx.append(cantSplit)
    
    return largura_total_tabela


# ============================================
# FUNÇÃO PRINCIPAL
# ============================================

def gerar_relatorio():
    """Função principal para gerar o relatório"""
    
    print("\n" + "="*60)
    print("📊 GERADOR DE RELATÓRIO DE METAS ESTRATÉGICAS - TJMG")
    print("="*60 + "\n")
    
    # 1. Criar pasta de saída
    criar_pasta_saida()
    
    # 2. Carregar dados
    df = carregar_dados()
    if df is None:
        return
    
    # 3. Carregar mapeamento de superintendências
    mapeamento = carregar_mapeamento_superintendencias()
    if not mapeamento:
        print("⚠️  Aviso: Mapeamento não carregado. Continuando sem classificação por superintendência.")
        return
    
    # 4. Adicionar coluna de superintendência ao DataFrame
    df = adicionar_coluna_superintendencia(df, mapeamento)
    
    # 5. Agrupar dados por superintendência e macrodesafio
    grupos_super = agrupar_por_superintendencia_e_macro(df)
    
    # 6. Criar documento com primeira página em retrato
    print("📝 Criando documento Word...")
    primeira_superintendencia = list(grupos_super.keys())[0] if grupos_super else 'Presidência'
    doc = criar_documento(primeira_superintendencia)
    
    # 7. Adicionar tabela histórica na primeira página (retrato)
    print("📊 Adicionando tabela histórica...")
    adicionar_cabecalho_relatorio(doc)
    
    # 8. Criar segunda seção em paisagem para as tabelas de metas
    print("📄 Criando seção em paisagem...")
    criar_secao_paisagem_inicial(doc, primeira_superintendencia)
    
    # 9. Gerar seções por superintendência no mesmo documento
    print("✍️  Gerando seções do relatório por Superintendência...")
    
    primeira_super = True
    for superintendencia, grupos_macro in grupos_super.items():
        print(f"\n📋 Superintendência: {superintendencia}")
        
        # Adicionar nova seção com cabeçalho específico (exceto para a primeira)
        if not primeira_super:
            adicionar_nova_secao_superintendencia(doc, superintendencia, primeira=False)
        
        # Adicionar cada Macrodesafio desta superintendência
        primeira_secao_super = True
        for idx, (macrodesafio, df_grupo) in enumerate(grupos_macro):
            print(f"   → {macrodesafio} ({len(df_grupo)} registros)")
            adicionar_secao_macrodesafio(doc, macrodesafio, df_grupo, primeira_secao=primeira_secao_super)
            primeira_secao_super = False
        
        primeira_super = False
    
    # 10. Salvar documento único
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_arquivo = f"{Config.NOME_RELATORIO}_{timestamp}.docx"
    caminho_completo = os.path.join(Config.PASTA_SAIDA, nome_arquivo)
    
    doc.save(caminho_completo)
    
    print(f"\n✅ RELATÓRIO GERADO COM SUCESSO!")
    print(f"📁 Localização: {os.path.abspath(caminho_completo)}")
    print(f"📄 Total de Superintendências: {len(grupos_super)}")
    print(f"📊 Total de registros: {len(df)}")
    print("\n" + "="*60 + "\n")


# ============================================
# EXECUÇÃO
# ============================================

if __name__ == "__main__":
    try:
        gerar_relatorio()
    except KeyboardInterrupt:
        print("\n\n⚠️  Geração cancelada pelo usuário.")
    except Exception as e:
        print(f"\n❌ ERRO INESPERADO: {e}")
        import traceback
        traceback.print_exc()